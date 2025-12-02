# -------------------- Routes --------------------

@app.route('/')
def index():
    return render_template('index.html') if os.path.exists('templates/index.html') else "<h1>Welcome - index</h1>"

@app.route('/teacher_login', methods=['GET', 'POST'])
def teacher_login():
    if request.method == 'POST':
        user = request.form.get('username')
        pwd = request.form.get('password')
        if TEACHER_CREDENTIALS.get(user) == pwd:
            session['teacher_logged_in'] = True
            session['teacher_username'] = user
            return redirect(url_for('teacher_dashboard'))
        flash('Invalid credentials', 'danger')
    return render_template('teacher_login.html') if os.path.exists('templates/teacher_login.html') else "<form method='post'>Username: <input name='username'/><br/>Password: <input name='password' type='password'/><button>Login</button></form>"

@app.route('/teacher_logout')
def teacher_logout():
    session.pop('teacher_logged_in', None)
    return redirect(url_for('index'))

@app.route('/teacher_dashboard')
def teacher_dashboard():
    if not session.get('teacher_logged_in'):
        return redirect(url_for('teacher_login'))
    return render_template('teacher_dashboard.html') if os.path.exists('templates/teacher_dashboard.html') else "<h1>Teacher Dashboard</h1>"

@app.route('/select_level', methods=['GET', 'POST'])
def select_level():
    levels = load_levels() or []
    if request.method == 'POST':
        try:
            selected_level = int(request.form.get('level', 1))
        except Exception:
            selected_level = 1
        session['selected_level'] = selected_level
        settings = load_json_file('data/game_settings.json', {}) or {}
        session['score'] = 0
        session['player_hp'] = settings.get('base_player_hp', 100)
        session['enemy_hp'] = settings.get('base_enemy_hp', 50)
        session['q_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['level_start_time'] = time.time()
        session['level_completed'] = False
        session['enemy_index'] = 0
        # auto_save_progress()  # Uncomment if you have this helper
        return redirect(url_for('game'))
    return render_template('select_level.html', levels=levels) if os.path.exists('templates/select_level.html') else f"<h1>Select Level</h1><p>{len(levels)} levels available</p>"

@app.route('/game')
def game():
    # Minimal game page
    return render_template('game.html') if os.path.exists('templates/game.html') else "<h1>Game</h1>"

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(path)
        return jsonify({'success': True, 'filename': filename})
    return jsonify({'error': 'File type not allowed'}), 400
# Full cleaned Flask app - production-ready structure
# - Complete application with routes and fixed helpers
# - Defensive programming: guards for missing files, optional SocketIO, safe AI handling
# - Note: This file focuses on clarity and safety. Add your project-specific templates and static files.

import os
import sys
import json
import os
import sys
import json
import time
import random
import threading
import requests
import difflib
from datetime import datetime
from functools import wraps
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_from_directory

# Optional socket support
try:
    from flask_socketio import SocketIO
except Exception:
    SocketIO = None

# Whoosh optional (kept for compatibility if present)
try:
    from whoosh.fields import Schema, TEXT, ID
    from whoosh import index
except Exception:
    Schema = None
    index = None

# Load config with safe defaults
try:
    from config import (TEACHER_CREDENTIALS, AI_PROVIDER, OPENAI_API_KEY, GEMINI_API_KEY,
                        OPENAI_MODEL, GEMINI_MODEL, AI_MODEL, UPLOAD_FOLDER, ALLOWED_EXTENSIONS,
                        MAX_CONTENT_LENGTH)
except Exception:
    TEACHER_CREDENTIALS = {'admin': 'password'}
    AI_PROVIDER = 'openai'
    OPENAI_API_KEY = ''
    GEMINI_API_KEY = ''
    OPENAI_MODEL = 'gpt-4o'
    GEMINI_MODEL = 'gemini'
    AI_MODEL = None
    UPLOAD_FOLDER = 'uploads'
    ALLOWED_EXTENSIONS = {'txt', 'md', 'pdf', 'docx'}
    MAX_CONTENT_LENGTH = 16 * 1024 * 1024

# App init
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET', 'unix_rpg_secret')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
os.makedirs('data', exist_ok=True)
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Optional SocketIO init
socketio = None
if SocketIO is not None:
    try:
        socketio = SocketIO(app, cors_allowed_origins='*')
    except Exception as e:
        print(f"SocketIO unavailable: {e}")
        socketio = None

# Files
LEADERBOARD_FILE = os.path.join('data', 'leaderboard.json')
GUEST_LEADERBOARD_FILE = os.path.join('data', 'guest_leaderboard.json')
ANALYTICS_FILE = os.path.join('data', 'analytics.json')
STUDENTS_FILE = os.path.join('data', 'students.json')
QUESTIONS_FILE = os.path.join('data', 'questions.json')
LEVELS_FILE = os.path.join('data', 'levels.json')
CHAPTERS_FILE = os.path.join('data', 'chapters.json')

# -------------------- Helpers --------------------

def load_json_file(path, default=None):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default
    except Exception as e:
        print(f"Error loading {path}: {e}")
        return default

def save_json_file(path, data):
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving {path}: {e}")

def load_chapters():
    return load_json_file(CHAPTERS_FILE, {"chapters": []})

def load_students():
    return load_json_file(STUDENTS_FILE, [])

def load_questions():
    return load_json_file(QUESTIONS_FILE, [])

def load_levels():
    return load_json_file(LEVELS_FILE, [])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    SocketIO = None

# Whoosh optional (kept for compatibility if present)
try:
    from whoosh.fields import Schema, TEXT, ID
    from whoosh import index
except Exception:
    Schema = None
    index = None

# Load config with safe defaults
try:
    from config import (TEACHER_CREDENTIALS, AI_PROVIDER, OPENAI_API_KEY, GEMINI_API_KEY,
                        OPENAI_MODEL, GEMINI_MODEL, AI_MODEL, UPLOAD_FOLDER, ALLOWED_EXTENSIONS,
                        MAX_CONTENT_LENGTH)
except Exception:
    TEACHER_CREDENTIALS = {'admin': 'password'}
    AI_PROVIDER = 'openai'
    OPENAI_API_KEY = ''
    GEMINI_API_KEY = ''
    OPENAI_MODEL = 'gpt-4o'
    GEMINI_MODEL = 'gemini'
    AI_MODEL = None
    UPLOAD_FOLDER = 'uploads'
    ALLOWED_EXTENSIONS = {'txt', 'md', 'pdf', 'docx'}
    MAX_CONTENT_LENGTH = 16 * 1024 * 1024

# App init
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET', 'unix_rpg_secret')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
os.makedirs('data', exist_ok=True)
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Optional SocketIO init
socketio = None
if SocketIO is not None:
    try:
        socketio = SocketIO(app, cors_allowed_origins='*')
    except Exception as e:
        print(f"SocketIO unavailable: {e}")
        socketio = None

# Files
LEADERBOARD_FILE = os.path.join('data', 'leaderboard.json')
GUEST_LEADERBOARD_FILE = os.path.join('data', 'guest_leaderboard.json')
ANALYTICS_FILE = os.path.join('data', 'analytics.json')
STUDENTS_FILE = os.path.join('data', 'students.json')
QUESTIONS_FILE = os.path.join('data', 'questions.json')
LEVELS_FILE = os.path.join('data', 'levels.json')
CHAPTERS_FILE = os.path.join('data', 'chapters.json')

# -------------------- Helpers --------------------

def load_json_file(path, default=None):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default
    except Exception as e:
        print(f"Error loading {path}: {e}")
        return default


def save_json_file(path, data):
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving {path}: {e}")


def load_chapters():
    return load_json_file(CHAPTERS_FILE, {"chapters": []})


def load_students():
    return load_json_file(STUDENTS_FILE, [])


def load_questions():
    return load_json_file(QUESTIONS_FILE, [])


def load_levels():
    return load_json_file(LEVELS_FILE, [])


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# -------------------- AI Helpers --------------------

def is_ai_configured():
    if AI_PROVIDER == 'gemini':
        return bool(GEMINI_API_KEY and GEMINI_API_KEY.strip() and len(GEMINI_API_KEY) > 10)
    elif AI_PROVIDER == 'openai':
        return bool(OPENAI_API_KEY and OPENAI_API_KEY.strip() and len(OPENAI_API_KEY) > 10)
    return False


def get_ai_config_error_message():
    if AI_PROVIDER == 'gemini':
        return 'Gemini API key not configured. Please set GEMINI_API_KEY in config.py.'
    elif AI_PROVIDER == 'openai':
        return 'OpenAI API key not configured. Please set OPENAI_API_KEY in config.py.'
    return 'AI provider not configured.'


def call_openai_api(prompt, max_tokens=1000):
    try:
        if not OPENAI_API_KEY:
            return 'OpenAI API key missing'
        headers = {'Authorization': f'Bearer {OPENAI_API_KEY}', 'Content-Type': 'application/json'}
        data = {'model': OPENAI_MODEL, 'messages': [{'role': 'user', 'content': prompt}], 'max_tokens': max_tokens}
        resp = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=data, timeout=60)
        if resp.status_code == 200:
            return resp.json()['choices'][0]['message']['content']
        return f"Error: {resp.status_code} - {resp.text}"
    except Exception as e:
        return f"Error calling OpenAI: {e}"


def call_gemini_api(prompt, max_tokens=1000):
    try:
        if not GEMINI_API_KEY:
            return 'Gemini API key missing'
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}'
        data = {'contents': [{'parts': [{'text': prompt}]}], 'generationConfig': {'maxOutputTokens': max_tokens}}
        resp = requests.post(url, json=data, timeout=60)
        if resp.status_code == 200:
            result = resp.json()
            if 'candidates' in result and len(result['candidates']) > 0:
                cand = result['candidates'][0]
                if 'content' in cand and 'parts' in cand['content']:
                    return cand['content']['parts'][0]['text']
                return 'Gemini: unexpected response format'
            return 'Gemini: no candidates returned'
        return f"Gemini Error: {resp.status_code} - {resp.text}"
    except Exception as e:
        return f"Error calling Gemini: {e}"


def call_ai_api(prompt, max_tokens=1000):
    if AI_PROVIDER == 'gemini':
        return call_gemini_api(prompt, max_tokens)
    return call_openai_api(prompt, max_tokens)

# -------------------- Safe emit --------------------

def safe_emit(event, data, room=None):
    if socketio:
        try:
            socketio.emit(event, data, room=room)
        except Exception as e:
            print(f"Socket emit failed: {e}")

# -------------------- Game Logic Helpers --------------------

def generate_enemy_taunt(question, enemy_name):
    if not isinstance(question, dict):
        return f"{enemy_name} challenges your knowledge!"
    qtext = str(question.get('q', '')).lower()
    keywords = question.get('keywords', [])
    if isinstance(keywords, str):
        keywords = [k.strip().lower() for k in keywords.split(',') if k.strip()]
    else:
        keywords = [str(k).strip().lower() for k in (keywords or [])]

    taunt_templates = {
        'ls': ["Can you even list a directory?", "Let's see if you know what 'ls' does!"],
        'cd': ["Lost in the filesystem already?", "Can you navigate directories?"],
        'chmod': ["Permissions confuse you, don't they?", "Can you handle file permissions?"],
        'grep': ["Can you search through text?", "Let's test your pattern matching!"],
        'network': ["Networking will be your downfall!", "Can you configure network settings?"],
    }

    for kw in keywords:
        for topic, taunts in taunt_templates.items():
            if topic in kw or topic in qtext:
                return random.choice(taunts)
    for topic, taunts in taunt_templates.items():
        if topic in qtext:
            return random.choice(taunts)
    return random.choice([f"{enemy_name} challenges your knowledge!", "Prove your expertise!"])


def apply_adaptive_difficulty(base_questions, level_number):
    questions = load_questions()
    if not isinstance(base_questions, list):
        return []

    correct = session.get('correct_answers', 0)
    wrong = session.get('wrong_answers', 0)
    total = correct + wrong
    if total == 0:
        return base_questions
    accuracy = correct / total

    if accuracy > 0.8:
        all_levels = load_levels() or []
        harder = []
        for lvl in all_levels:
            if lvl.get('level', 0) > level_number:
                harder.extend([q for q in questions if q.get('id') in lvl.get('questions', [])])
        if harder:
            base_count = int(len(base_questions) * 0.7)
            hard_count = min(len(harder), max(0, len(base_questions) - base_count))
            return base_questions[:base_count] + random.sample(harder, hard_count)

    if accuracy < 0.5:
        all_levels = load_levels() or []
        easier = []
        for lvl in all_levels:
            if lvl.get('level', 0) < level_number:
                easier.extend([q for q in questions if q.get('id') in lvl.get('questions', [])])
        if easier:
            base_count = int(len(base_questions) * 0.7)
            easy_count = min(len(easier), max(0, len(base_questions) - base_count))
            return base_questions[:base_count] + random.sample(easier, easy_count)

    return base_questions

# -------------------- Fuzzy grading --------------------

def normalize_true_false_answer(answer):
    a = str(answer).lower().strip()
    if a in ['true', 't', 'yes', 'y', '1', 'correct', 'right']:
        return 'true'
    if a in ['false', 'f', 'no', 'n', '0', 'incorrect', 'wrong']:
        return 'false'
    return a


def check_answer_fuzzy(user_answer, question_data, similarity_threshold=0.8):
    ua = str(user_answer).strip()
    qtype = question_data.get('type', 'short_answer')
    ca = str(question_data.get('answer', '')).strip()

    if qtype == 'true_false':
        return normalize_true_false_answer(ua) == normalize_true_false_answer(ca), 'true_false', 1.0

    if qtype == 'multiple_choice':
        ua_low = ua.lower()
        ca_low = ca.lower()
        options = question_data.get('options', [])
        if ua_low == ca_low:
            return True, 'Correct choice!', 1.0
        if len(ua_low) == 1 and ua_low in 'abcd':
            idx = ord(ua_low) - ord('a')
            if 0 <= idx < len(options) and options[idx].lower().strip() == ca_low:
                return True, 'Correct choice!', 1.0
        for opt in options:
            if opt.lower().strip() == ua_low and opt.lower().strip() == ca_low:
                return True, 'Correct choice!', 1.0
        return False, 'Incorrect choice', 0

    ua_low = ua.lower()
    ca_low = ca.lower()
    if ua_low == ca_low:
        return True, 'Exact match!', 1.0

    keywords = question_data.get('keywords', [])
    if isinstance(keywords, str):
        keywords = [k.strip().lower() for k in keywords.split(',') if k.strip()]
    else:
        keywords = [str(k).strip().lower() for k in (keywords or [])]

    if ua_low in keywords:
        return True, 'Correct alternative!', 1.0

    if len(ua_low) > 3 and (ua_low in ca_low or ca_low in ua_low):
        return True, 'Partial match!', 0.8

    sim = difflib.SequenceMatcher(None, ua_low, ca_low).ratio()
    if sim > similarity_threshold:
        return True, 'Close match', sim

    best_sim = 0
    best_kw = ''
    for kw in keywords:
        s = difflib.SequenceMatcher(None, ua_low, kw).ratio()
        if s > best_sim:
            best_sim = s
            best_kw = kw
    if best_sim > similarity_threshold:
        return True, f"Close enough to '{best_kw}'", best_sim

    uw = set(ua_low.split())
    cw = set(ca_low.split())
    if uw and cw:
        overlap = len(uw.intersection(cw)) / max(1, len(cw))
        if overlap >= 0.7:
            return True, 'Word-based match', overlap

    return False, 'Incorrect', 0

# -------------------- Leaderboards --------------------

def save_leaderboard(player_name, score, total_time, correct_answers, wrong_answers, game_mode='adventure', level=None):
    is_student = session.get('is_student') and session.get('student_id')
    record = {
        'player': player_name,
        'score': score,
        'time': round(total_time, 2),
        'correct_answers': correct_answers,
        'wrong_answers': wrong_answers,
        'game_mode': game_mode,
        'date': time.strftime('%Y-%m-%d %H:%M:%S')
    }
    if level is not None:
        record['level'] = level

    try:
        if is_student:
            students = load_students()
            sid = session.get('student_id')
            student = next((s for s in students if s.get('id') == sid), None)
            if student and student.get('full_name'):
                record['player'] = student['full_name']
            lb = load_json_file(LEADERBOARD_FILE, []) or []
            lb.append(record)
            save_json_file(LEADERBOARD_FILE, lb)
        else:
            glb = load_json_file(GUEST_LEADERBOARD_FILE, []) or []
            glb.append(record)
            save_json_file(GUEST_LEADERBOARD_FILE, glb)
    except Exception as e:
        print(f"Error saving leaderboard: {e}")

# -------------------- Analytics --------------------

def log_analytics_event(event_type, data=None):
    settings = load_json_file('data/game_settings.json', {}) or {}
    if not settings.get('analytics_enabled', True):
        return
    try:
        analytics = load_json_file(ANALYTICS_FILE, []) or []
        entry = {'timestamp': time.time(), 'event_type': event_type, 'player_name': session.get('player_name', 'Anonymous'), 'session_id': session.get('_id', 'unknown'), 'data': data or {}}
        analytics.append(entry)
        if len(analytics) > 1000:
            analytics = analytics[-1000:]
        save_json_file(ANALYTICS_FILE, analytics)
    except Exception as e:
        print(f"Analytics logging failed: {e}")

# -------------------- Routes --------------------

@app.route('/')
def index():
    return render_template('index.html') if os.path.exists('templates/index.html') else "<h1>Welcome - index</h1>"


@app.route('/teacher_login', methods=['GET', 'POST'])
def teacher_login():
    if request.method == 'POST':
        user = request.form.get('username')
        pwd = request.form.get('password')
        if TEACHER_CREDENTIALS.get(user) == pwd:
            session['teacher_logged_in'] = True
            session['teacher_username'] = user
            return redirect(url_for('teacher_dashboard'))
        flash('Invalid credentials', 'danger')
    return render_template('teacher_login.html') if os.path.exists('templates/teacher_login.html') else "<form method='post'>Username: <input name='username'/><br/>Password: <input name='password' type='password'/><button>Login</button></form>"


@app.route('/teacher_logout')
def teacher_logout():
    session.pop('teacher_logged_in', None)
    return redirect(url_for('index'))


@app.route('/teacher_dashboard')
def teacher_dashboard():
    if not session.get('teacher_logged_in'):
        return redirect(url_for('teacher_login'))
    return render_template('teacher_dashboard.html') if os.path.exists('templates/teacher_dashboard.html') else "<h1>Teacher Dashboard</h1>"


@app.route('/select_level', methods=['GET', 'POST'])
def select_level():
    levels = load_levels() or []
    if request.method == 'POST':
        try:
            selected_level = int(request.form.get('level', 1))
        except Exception:
            selected_level = 1
        session['selected_level'] = selected_level
        settings = load_json_file('data/game_settings.json', {}) or {}
        session['score'] = 0
        session['player_hp'] = settings.get('base_player_hp', 100)
        session['enemy_hp'] = settings.get('base_enemy_hp', 50)
        session['q_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['level_start_time'] = time.time()
        session['level_completed'] = False
        session['enemy_index'] = 0
        auto_save_progress()
        return redirect(url_for('game'))
    return render_template('select_level.html', levels=levels) if os.path.exists('templates/select_level.html') else f"<h1>Select Level</h1><p>{len(levels)} levels available</p>"


@app.route('/game')
def game():
    # Minimal game page
    return render_template('game.html') if os.path.exists('templates/game.html') else "<h1>Game</h1>"


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(path)
        return jsonify({'success': True, 'filename': filename})
    return jsonify({'error': 'File type not allowed'}), 400


@app.route('/download/<path:filename>')
def download_file(filename):
    # Secure file download
    safe_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(safe_path):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    return "File not found", 404


# ------------------- ENDLESS MODE -------------------
# (Moved below app initialization)
# This code ensures Flask and Whoosh are installed before importing them, preventing runtime errors
import subprocess
import sys
import os
import json
import time
import requests
import hashlib
import difflib
from datetime import datetime
from functools import wraps
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from flask_socketio import SocketIO, emit, join_room, leave_room
from whoosh.fields import Schema, TEXT, ID
from whoosh import index
from config import TEACHER_CREDENTIALS, AI_PROVIDER, OPENAI_API_KEY, GEMINI_API_KEY, OPENAI_MODEL, GEMINI_MODEL, AI_MODEL, UPLOAD_FOLDER, ALLOWED_EXTENSIONS, MAX_CONTENT_LENGTH
import threading

# Constants
LEADERBOARD_FILE = "data/leaderboard.json"
GUEST_LEADERBOARD_FILE = "data/guest_leaderboard.json"

# Load game settings at startup
def load_initial_settings():
    """Load initial game settings or use defaults"""
    try:
        settings_file = 'data/game_settings.json'
        with open(settings_file, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        return settings
    except (FileNotFoundError, json.JSONDecodeError):
        # Create default settings file if it doesn't exist
        default_settings = {
            'base_player_hp': 100,
            'base_enemy_hp': 50,
            'base_damage': 10,
            'question_time_limit': 30,
            'questions_per_level': 10
        }
        try:
            os.makedirs('data', exist_ok=True)
            with open('data/game_settings.json', 'w', encoding='utf-8') as f:
                json.dump(default_settings, f, indent=2)
        except Exception:
            pass
        return default_settings

# Initialize settings
initial_settings = load_initial_settings()
BASE_DAMAGE = initial_settings.get('base_damage', 10)
BASE_ENEMY_HP = initial_settings.get('base_enemy_hp', 50)
LEVEL_TIME_LIMIT = initial_settings.get('question_time_limit', 30)

# Function to generate dynamic enemy taunts based on question
def generate_enemy_taunt(question, enemy_name):
    """Generate a dynamic taunt based on the question content"""
    question_text = question.get('q', '').lower()
    keywords = question.get('keywords', [])
    
    # Extract key topics from question
    if isinstance(keywords, str):
        keywords = [k.strip().lower() for k in keywords.split(',')]

    taunt_templates = {
        'ls': [
            "Can you even list a directory?",
            "Let's see if you know what 'ls' does!",
            "Basic commands? This should be easy... or not!"
        ],
        'cd': [
            "Lost in the filesystem already?",
            "Can you navigate directories?",
            "Where do you think you're going?"
        ],
        'chmod': [
            "Permissions confuse you, don't they?",
            "Can you handle file permissions?",
            "Let's test your permission knowledge!"
        ],
        'systemctl': [
            "Service management is my domain!",
            "Can you control system services?",
            "Let's see your systemctl skills!"
        ],
        'firewall': [
            "Your firewall knowledge is weak!",
            "Can you protect this system?",
            "Let's test your security skills!"
        ],
        'user': [
            "User management is tricky, isn't it?",
            "Can you handle users and groups?",
            "Let's see if you can manage users!"
        ],
        'network': [
            "Networking will be your downfall!",
            "Can you configure network settings?",
            "Let's test your network knowledge!"
        ],
        'mount': [
            "Can you mount filesystems?",
            "Storage management is complex!",
            "Let's see your mounting skills!"
        ],
        'selinux': [
            "SELinux is too advanced for you!",
            "Can you handle security contexts?",
            "Security-Enhanced Linux will defeat you!"
        ],
        'lvm': [
            "Logical volumes will confuse you!",
            "Can you manage LVM?",
            "Storage management is my specialty!"
        ],
        'cron': [
            "Can you schedule tasks?",
            "Time-based jobs are tricky!",
            "Let's test your automation skills!"
        ],
        'package': [
            "Package management is complex!",
            "Can you install software?",
            "Let's see your package skills!"
        ],
        'grep': [
            "Can you search through text?",
            "Let's test your pattern matching!",
            "Grep will be your challenge!"
        ],
        'find': [
            "Can you find files?",
            "Let's see your search skills!",
            "File searching will defeat you!"
        ],
        'tar': [
            "Archiving is too complex for you!",
            "Can you handle tar archives?",
            "Let's test your compression knowledge!"
        ],
        'dns': [
            "DNS will confuse you!",
            "Can you resolve hostnames?",
            "Let's test your name resolution skills!"
        ],
        'boot': [
            "Boot processes are tricky!",
            "Can you fix boot issues?",
            "System startup will challenge you!"
        ]
    }
    
    # Find matching topic
    for keyword in keywords:
        for topic, taunts in taunt_templates.items():
            if topic in keyword or topic in question_text:
                import random
                return random.choice(taunts)
    
    # Check question text for topics
    for topic, taunts in taunt_templates.items():
        if topic in question_text:
            import random
            return random.choice(taunts)
    
    # Generic taunts based on enemy name
    generic_taunts = [
        f"{enemy_name} challenges your knowledge!",
        "This question will test your skills!",
        "Can you answer this correctly?",
        "Let's see what you know!",
        "Prove your expertise!"
    ]
    
    import random
    return random.choice(generic_taunts)

# Define the schema for the Whoosh search index
schema = Schema(
    id=ID(stored=True, unique=True),
    question=TEXT(stored=True),
    answer=TEXT(stored=True),
    keywords=TEXT(stored=True)
)

# Create or open the Whoosh index directory
def create_or_open_index():
    """Open the Whoosh index if valid; otherwise recreate it."""
    try:
        if not os.path.exists("indexdir"):
            os.mkdir("indexdir")
            return index.create_in("indexdir", schema)
        try:
            return index.open_dir("indexdir")
        except Exception as e:
            print(f"Whoosh index open failed: {e}. Recreating indexdir...")
            try:
                # Remove the corrupted indexdir and recreate
                import shutil
                shutil.rmtree("indexdir")
            except Exception:
                pass
            os.mkdir("indexdir")
            return index.create_in("indexdir", schema)
    except Exception as e:
        print(f"Unexpected error creating/opening indexdir: {e}")
        raise

ix = create_or_open_index()

# Initialize the Flask app
app = Flask(__name__)
import copy

app.secret_key = "unix_rpg_secret"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Initialize Flask-SocketIO
try:
    socketio = SocketIO(app, cors_allowed_origins="*")
except Exception as e:
    print(f"Warning: SocketIO initialization failed: {e}")
    socketio = None

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.before_request
def check_session_timeout():
    """Check if session has timed out based on settings"""
    settings = get_current_game_settings()
    session_timeout_minutes = settings.get('session_timeout', 30)
    
    # Skip timeout check for certain routes
    excluded_routes = ['static', 'teacher_login', 'teacher_logout']
    if request.endpoint in excluded_routes:
        return
    
    # Check if user has activity timestamp
    if 'last_activity' in session:
        last_activity = session['last_activity']
        current_time = time.time()
        timeout_seconds = session_timeout_minutes * 60
        
        if current_time - last_activity > timeout_seconds:
            # Session timed out - clear session and redirect
            session.clear()
            flash(f'Your session timed out after {session_timeout_minutes} minutes of inactivity.', 'warning')
            return redirect(url_for('index'))
    
    # Update last activity timestamp
    session['last_activity'] = time.time()

# Teacher authentication decorator
def teacher_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('teacher_logged_in'):
            return redirect(url_for('teacher_login'))
        return f(*args, **kwargs)
    return decorated_function

# Student authentication decorator
def student_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('is_student'):
            return redirect(url_for('student_login'))
        return f(*args, **kwargs)
    return decorated_function

# Helper function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# AI Integration Functions
def extract_json_from_response(response):
    """Extract JSON from AI response, handling markdown code blocks and truncation"""
    import re
    
    # If the response contains markdown code blocks, extract the JSON
    json_pattern = r'```(?:json)?\s*(\[.*\]|\{.*\})\s*```'
    match = re.search(json_pattern, response, re.DOTALL)
    
    if match:
        return match.group(1).strip()
    
    # For JSON arrays, find the opening [ and matching closing ]
    start_idx = response.find('[')
    if start_idx != -1:
        bracket_count = 0
        for i, char in enumerate(response[start_idx:], start_idx):
            if char == '[':
                bracket_count += 1
            elif char == ']':
                bracket_count -= 1
                if bracket_count == 0:
                    return response[start_idx:i+1].strip()
    
    # For JSON objects, find the opening { and matching closing }
    start_idx = response.find('{')
    if start_idx != -1:
        brace_count = 0
        for i, char in enumerate(response[start_idx:], start_idx):
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    return response[start_idx:i+1].strip()
    
    # If no JSON structure found, return the response as-is
    return response.strip()

def fix_incomplete_json(json_str):
    """Attempt to fix incomplete JSON by adding missing closing brackets/braces"""
    json_str = json_str.strip()
    
    # Count opening and closing brackets/braces
    open_brackets = json_str.count('[')
    close_brackets = json_str.count(']')
    open_braces = json_str.count('{')
    close_braces = json_str.count('}')
    
    # Add missing closing brackets
    while close_brackets < open_brackets:
        json_str += ']'
        close_brackets += 1
    
    # Add missing closing braces
    while close_braces < open_braces:
        json_str += '}'
        close_braces += 1
    
    # If the JSON ends with a comma, try to remove it
    if json_str.rstrip().endswith(','):
        json_str = json_str.rstrip()[:-1]
    
    return json_str

def call_ai_api(prompt, max_tokens=1000):
    """Call AI API (OpenAI or Gemini) with error handling"""
    if AI_PROVIDER == "gemini":
        return call_gemini_api(prompt, max_tokens)
    elif AI_PROVIDER == "openai":
        return call_openai_api(prompt, max_tokens)
    else:
        return {"error": "Invalid AI provider configured"}

def is_ai_configured():
    """Check if AI API key is properly configured"""
    if AI_PROVIDER == "gemini":
        return bool(GEMINI_API_KEY and GEMINI_API_KEY.strip() and GEMINI_API_KEY != 'your-gemini-api-key-here' and len(GEMINI_API_KEY) > 10)
    elif AI_PROVIDER == "openai":
        return bool(OPENAI_API_KEY and OPENAI_API_KEY.strip() and OPENAI_API_KEY != 'your-openai-api-key-here' and len(OPENAI_API_KEY) > 10)
    return False

def get_ai_config_error_message():
    """Get user-friendly error message for missing AI configuration"""
    if AI_PROVIDER == "gemini":
        return "Gemini API key not configured. Please set GEMINI_API_KEY in config.py. Get a free key at https://makersuite.google.com/app/apikey"
    elif AI_PROVIDER == "openai":
        return "OpenAI API key not configured. Please set OPENAI_API_KEY in config.py. Get a key at https://platform.openai.com/api-keys"
    return "AI provider not configured properly in config.py"

def call_openai_api(prompt, max_tokens=1000):
    """Call OpenAI API with error handling"""
    try:
        headers = {
            'Authorization': f'Bearer {OPENAI_API_KEY}',
            'Content-Type': 'application/json',
        }
        data = {
            'model': OPENAI_MODEL,
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': max_tokens,
            'temperature': 0.7
        }
        response = requests.post('https://api.openai.com/v1/chat/completions', 
                               headers=headers, json=data, timeout=60)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error calling OpenAI API: {str(e)}"

def call_gemini_api(prompt, max_tokens=1000):
    """Call Google Gemini API with error handling"""
    try:
        if not GEMINI_API_KEY or GEMINI_API_KEY == 'your-gemini-api-key-here':
            return "Gemini API Error: API key not configured. Please set GEMINI_API_KEY in config.py"
        
        headers = {
            'Content-Type': 'application/json'
        }
        
        # Gemini API endpoint
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}'
        
        data = {
            'contents': [{
                'parts': [{'text': prompt}]
            }],
            'generationConfig': {
                'maxOutputTokens': max_tokens,
                'temperature': 0.7
            }
        }
        
        print(f"DEBUG: Calling Gemini API with model: {GEMINI_MODEL}")
        print(f"DEBUG: API Key configured: {bool(GEMINI_API_KEY and len(GEMINI_API_KEY) > 20)}")
        print(f"DEBUG: Max tokens requested: {max_tokens}")
        
        response = requests.post(url, headers=headers, json=data, timeout=60)
        
        print(f"DEBUG: Gemini API response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"DEBUG: Full Gemini response structure: {json.dumps(result, indent=2)[:500]}")
            
            if 'candidates' in result and len(result['candidates']) > 0:
                candidate = result['candidates'][0]
                
                # Check if response was blocked
                if 'finishReason' in candidate and candidate['finishReason'] != 'STOP':
                    return f"Gemini API Error: Response blocked or incomplete (reason: {candidate.get('finishReason', 'unknown')})"
                
                if 'content' in candidate and 'parts' in candidate['content']:
                    response_text = candidate['content']['parts'][0]['text']
                    print(f"DEBUG: Response length: {len(response_text)} characters")
                    print(f"DEBUG: Response preview: {response_text[:200]}")
                    return response_text
                else:
                    print(f"DEBUG: Unexpected candidate structure: {candidate}")
                    return "Gemini API Error: Response missing content or parts"
            else:
                print(f"DEBUG: No candidates in response: {result}")
                return f"Gemini API Error: No response generated. Full response: {json.dumps(result)[:300]}"
        else:
            error_details = response.text
            print(f"DEBUG: Gemini API Error Details: {error_details}")
            
            if response.status_code == 400:
                return "Gemini API Error: Invalid API key or request format. Please check your API key and try again."
            elif response.status_code == 403:
                return "Gemini API Error: API key doesn't have permission or quota exceeded."
            elif response.status_code == 429:
                return "Gemini API Error: Rate limit exceeded. Please try again later."
            else:
                return f"Gemini API Error: {response.status_code} - {error_details}"
            
    except Exception as e:
        print(f"DEBUG: Exception in call_gemini_api: {str(e)}")
        return f"Error calling Gemini API: {str(e)}"

def extract_text_from_file(file_path):
    """Extract text content from uploaded files"""
    try:
        if file_path.endswith('.txt') or file_path.endswith('.md'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.endswith('.pdf'):
            # You would need to install PyPDF2: pip install PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text
            except ImportError:
                return "PDF support requires PyPDF2. Install with: pip install PyPDF2"
        elif file_path.endswith('.docx'):
            # You would need to install python-docx: pip install python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                return "DOCX support requires python-docx. Install with: pip install python-docx"
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading file: {str(e)}"

def generate_questions_with_ai(content, topic, difficulty, question_count, context="", question_types=None):
    """Generate questions using AI based on content"""
    
    # Check if AI is configured
    if not is_ai_configured():
        return {"error": get_ai_config_error_message()}
    
    # Default question types if none specified
    if question_types is None:
        question_types = ["short_answer", "multiple_choice", "true_false"]
    
    # Calculate appropriate content length based on question count
    # More questions need more tokens for response, so reduce content accordingly
    base_content_limit = 6000
    content_limit = max(2000, base_content_limit - (question_count * 200))
    
    # Truncate content if needed
    if len(content) > content_limit:
        content = content[:content_limit] + "\n[Content truncated for AI processing...]"
    
    question_types_str = ", ".join(question_types)
    
    prompt = f"""Based on the following educational content about {topic}, generate EXACTLY {question_count} {difficulty}-level questions.

Content:
{content}

Additional Context: {context}

Generate a mix of question types from: {question_types_str}

Generate questions in this EXACT JSON format (no deviations):
[
    {{
        "q": "Question text here?",
        "answer": "correct answer",
        "keywords": ["alternative1", "alternative2"],
        "feedback": "Educational explanation",
        "type": "short_answer|multiple_choice|true_false",
        "options": ["option1", "option2", "option3", "option4"]
    }}
]

CRITICAL REQUIREMENTS:
- Return EXACTLY {question_count} questions in the array
- Each question must have fields: q, answer, keywords, feedback, type
- For multiple_choice questions: include "options" array with 2-4 choices, answer must match one option exactly
- For true_false questions: answer must be "true" or "false", no options needed
- For short_answer questions: include keywords array for alternative answers, no options needed
- Keep answers concise (under 50 words)
- Keep feedback brief (under 100 words)
- Focus on {difficulty} difficulty level
- Questions should test understanding of {topic}
- Return ONLY the JSON array, no other text
- Do NOT use markdown code blocks
- Ensure valid JSON with proper commas and brackets

Start response with [ and end with ]"""
    
    response = call_ai_api(prompt, max_tokens=5000)
    try:
        # Extract JSON from response (handles markdown code blocks)
        clean_json = extract_json_from_response(response)
        
        # Try to fix incomplete JSON
        fixed_json = fix_incomplete_json(clean_json)
        
        print(f"DEBUG: Original response length: {len(response)}")
        print(f"DEBUG: Extracted JSON length: {len(clean_json)}")
        print(f"DEBUG: Fixed JSON: {fixed_json[:200]}...")
        
        questions_data = json.loads(fixed_json)
        return questions_data
    except json.JSONDecodeError as e:
        # If not valid JSON, return error with more details
        print(f"DEBUG: JSON Parse Error at position {e.pos}: {str(e)}")
        print(f"DEBUG: Problematic JSON section: {response[max(0, e.pos-50):e.pos+50]}")
        return {"error": f"AI returned invalid JSON: {response[:500]}... | Parse error: {str(e)}"}

def grade_answer_with_ai(question, correct_answer, student_answer, confidence_threshold=80):
    """Use AI to grade student answers with semantic understanding"""
    # Check if AI is configured
    if not is_ai_configured():
        return {
            "correct": False,
            "confidence": 0,
            "explanation": "AI grading not available - API key not configured"
        }
    
    prompt = f"""
    Grade this student answer using semantic understanding:
    
    Question: {question}
    Expected Answer: {correct_answer}
    Student Answer: {student_answer}
    
    Evaluate if the student answer is semantically correct, even if worded differently.
    Consider synonyms, alternative phrasings, and equivalent commands/concepts.
    
    Respond in this JSON format:
    {{
        "correct": true/false,
        "confidence": 0-100,
        "explanation": "Brief explanation of your decision"
    }}
    
    Be generous with partial credit for answers that show understanding.
    
    IMPORTANT: Return ONLY the JSON object. Do not use markdown code blocks. Do not include any explanatory text. Just the raw JSON object.
    """
    
    response = call_ai_api(prompt, max_tokens=200)
    
    # Check if response is an error message
    if isinstance(response, str) and ('error' in response.lower() or 'api' in response.lower()):
        return {
            "correct": False,
            "confidence": 0,
            "explanation": f"AI grading error: {response[:100]}"
        }
    
    try:
        # Extract JSON from response (handles markdown code blocks)
        clean_json = extract_json_from_response(response)
        result = json.loads(clean_json)
        # Apply confidence threshold
        if result.get('confidence', 0) < confidence_threshold:
            result['correct'] = False
            result['explanation'] += f" (Below {confidence_threshold}% confidence threshold)"
        return result
    except json.JSONDecodeError as e:
        return {
            "correct": False,
            "confidence": 0,
            "explanation": "Error: Could not parse AI response"
        }

# Helper function for fuzzy answer checking
def check_answer_fuzzy(user_answer, question_data, similarity_threshold=0.8):
    """
    Check user answer against correct answer and alternatives with fuzzy matching.
    Returns tuple: (is_correct, feedback_type, similarity_score)
    """
    question_type = question_data.get('type', 'short_answer')
    user_answer = user_answer.strip()
    correct_answer = question_data.get("answer", "").strip()
    
    # Handle different question types
    if question_type == 'true_false':
        # Normalize true/false answers
        user_normalized = normalize_true_false_answer(user_answer)
        correct_normalized = normalize_true_false_answer(correct_answer)
        
        if user_normalized == correct_normalized:
            return True, "Correct!", 1.0
        else:
            return False, "Incorrect", 0
    
    elif question_type == 'multiple_choice':
        # For multiple choice, check exact match with correct answer or option labels
        user_answer_lower = user_answer.lower().strip()
        correct_answer_lower = correct_answer.lower().strip()
        
        # Check exact match
        if user_answer_lower == correct_answer_lower:
            return True, "Correct choice!", 1.0
        
        # Check if user entered option letter (a, b, c, d)
        options = question_data.get('options', [])
        if len(user_answer) == 1 and user_answer.lower() in 'abcd':
            option_index = ord(user_answer.lower()) - ord('a')
            if 0 <= option_index < len(options):
                if options[option_index].lower().strip() == correct_answer_lower:
                    return True, "Correct choice!", 1.0
        
        # Check if user typed the full option text
        for option in options:
            if option.lower().strip() == user_answer_lower:
                if option.lower().strip() == correct_answer_lower:
                    return True, "Correct choice!", 1.0
        
        return False, "Incorrect choice", 0
    
    else:  # short_answer (default) - use enhanced fuzzy matching
        user_answer = user_answer.lower().strip()
        correct_answer = correct_answer.lower().strip()
        
        # Check exact match with correct answer
        if user_answer == correct_answer:
            return True, "Exact match!", 1.0
        
        # Check exact match with keywords (alternatives)
        keywords = question_data.get("keywords", [])
        if isinstance(keywords, str):
            keywords = [k.strip().lower() for k in keywords.split(',') if k.strip()]
        else:
            keywords = [str(k).strip().lower() for k in keywords]
        
        if user_answer in keywords:
            return True, "Correct alternative!", 1.0
        
        # Enhanced fuzzy matching for short answers
        # 1. Check for partial matches (substring)
        if len(user_answer) > 3:  # Only for longer answers
            if user_answer in correct_answer or correct_answer in user_answer:
                return True, "Partial match!", 0.8
            
            # Check partial matches with keywords
            for keyword in keywords:
                if user_answer in keyword or keyword in user_answer:
                    return True, f"Partial match with '{keyword}'!", 0.8
        
        # 2. Fuzzy matching with correct answer
        correct_similarity = difflib.SequenceMatcher(None, user_answer, correct_answer).ratio()
        if correct_similarity > similarity_threshold:
            return True, "Close enough to correct answer!", correct_similarity
        
        # 3. Fuzzy matching with alternatives
        best_similarity = 0
        best_match = ""
        for keyword in keywords:
            similarity = difflib.SequenceMatcher(None, user_answer, keyword).ratio()
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = keyword
        
        if best_similarity > similarity_threshold:
            return True, f"Close enough to '{best_match}'!", best_similarity
        
        # 4. Word-based fuzzy matching (split into words)
        user_words = set(user_answer.split())
        correct_words = set(correct_answer.split())
        
        if user_words and correct_words:
            word_overlap = len(user_words.intersection(correct_words)) / len(correct_words)
            if word_overlap >= 0.7:  # 70% word overlap
                return True, "Word-based match!", word_overlap
        
        return False, "Incorrect", 0

def normalize_true_false_answer(answer):
    """Normalize true/false answers to handle various inputs"""
    answer_lower = answer.lower().strip()
    
    # True variations
    if answer_lower in ['true', 't', 'yes', 'y', '1', 'correct', 'right']:
        return 'true'
    
    # False variations
    if answer_lower in ['false', 'f', 'no', 'n', '0', 'incorrect', 'wrong']:
        return 'false'
    
    return answer_lower

# Helper function to save leaderboard data
def save_leaderboard(player_name, score, total_time, correct_answers, wrong_answers, game_mode="adventure", level=None):
    print(f"[DEBUG SAVE_LEADERBOARD] Called with: player={player_name}, score={score}, mode={game_mode}, is_student={session.get('is_student')}")
    
    # Save to student leaderboard if it's a logged-in student
    if session.get('is_student') and session.get('student_id'):
        # Get actual student name from students.json instead of relying on player_name
        student_id = session.get('student_id')
        students = load_students()
        student = next((s for s in students if s['id'] == student_id), None)
        
        # Use student's full name if available, otherwise use the passed player_name
        if student and 'full_name' in student:
            player_name = student['full_name']
        elif session.get('student_name'):
            player_name = session.get('student_name')
        
        print(f"Saving student leaderboard data: {player_name}, mode: {game_mode}, level: {level}")
        
        record = {
            "player": player_name,
            "student_id": student_id,
            "score": score,
            "time": round(total_time, 2),
            "correct_answers": correct_answers,
            "wrong_answers": wrong_answers,
            "game_mode": game_mode,
            "date": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # Add level for adventure mode
        if game_mode == "adventure" and level is not None:
            record["level"] = level
        
        try:
            with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
                leaderboard = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            leaderboard = []

        leaderboard.append(record)

        with open(LEADERBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(leaderboard, f, indent=4)
        
        print(f"[DEBUG SAVE_LEADERBOARD] Successfully saved to STUDENT leaderboard: {LEADERBOARD_FILE}")
    
    # Save to guest leaderboard if it's a guest player (not a student)
    else:
        print(f"Saving guest leaderboard data: {player_name}, mode: {game_mode}, level: {level}")
        
        record = {
            "player": player_name,
            "score": score,
            "time": round(total_time, 2),
            "correct_answers": correct_answers,
            "wrong_answers": wrong_answers,
            "game_mode": game_mode,
            "date": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # Add level for adventure mode
        if game_mode == "adventure" and level is not None:
            record["level"] = level
        
        try:
            with open(GUEST_LEADERBOARD_FILE, "r", encoding="utf-8") as f:
                guest_leaderboard = json.load(f)
            print(f"[DEBUG SAVE_LEADERBOARD] Loaded {len(guest_leaderboard)} existing guest records")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"[DEBUG SAVE_LEADERBOARD] Creating new guest leaderboard file: {e}")
            guest_leaderboard = []

        guest_leaderboard.append(record)
        print(f"[DEBUG SAVE_LEADERBOARD] Added record, now have {len(guest_leaderboard)} total records")

        with open(GUEST_LEADERBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(guest_leaderboard, f, indent=4)
        
        print(f"[DEBUG SAVE_LEADERBOARD] Successfully saved to GUEST leaderboard: {GUEST_LEADERBOARD_FILE}")

def load_leaderboard():
    """Load leaderboard data from file"""
    try:
        with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def reset_test_yourself_session():
    """Completely reset Test Yourself mode session data"""
    test_keys = ['test_question_ids', 'test_q_index', 'test_correct', 
                'test_start_time', 'test_time_limit', 'test_user_answers']
    for key in test_keys:
        session.pop(key, None)

def reset_endless_mode_session():
    """Completely reset Endless mode session data"""
    endless_keys = ['endless_score', 'endless_hp', 'endless_streak', 'endless_highest_streak',
                   'endless_total_answered', 'endless_correct', 'endless_wrong', 'endless_start_time',
                   'endless_score_initialized', 'endless_question_start', 'endless_current_question',
                   'endless_feedback_list']
    for key in endless_keys:
        session.pop(key, None)

# Auto-save functionality
def log_analytics_event(event_type, data=None):
    """Log analytics event if analytics are enabled"""
    settings = get_current_game_settings()
    if not settings.get('analytics_enabled', True):
        return
    
    try:
        analytics_data = {
            'timestamp': time.time(),
            'event_type': event_type,
            'player_name': session.get('player_name', 'Anonymous'),
            'session_id': session.get('_id', 'unknown'),
            'data': data or {}
        }
        
        # Append to analytics file
        analytics_file = 'data/analytics.json'
        analytics_log = []
        
        # Load existing analytics
        try:
            with open(analytics_file, 'r', encoding='utf-8') as f:
                analytics_log = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            analytics_log = []
        
        analytics_log.append(analytics_data)
        
        # Keep only last 1000 events to prevent file from growing too large
        if len(analytics_log) > 1000:
            analytics_log = analytics_log[-1000:]
        
        # Save analytics
        os.makedirs('data', exist_ok=True)
        with open(analytics_file, 'w', encoding='utf-8') as f:
            json.dump(analytics_log, f, indent=2)
            
    except Exception as e:
        print(f"Analytics logging failed: {e}")

def log_student_answer(student_id, student_name, question_id, question_text, student_answer, correct_answer, is_correct, game_mode, level=None):
    """Log student answers in real-time and notify teachers via WebSocket"""
    try:
        # Create answer log entry
        answer_log = {
            'timestamp': datetime.now().isoformat(),
            'student_id': student_id,
            'student_name': student_name,
            'question_id': question_id,
            'question_text': question_text[:100] + '...' if len(question_text) > 100 else question_text,
            'student_answer': student_answer,
            'correct_answer': correct_answer,
            'is_correct': is_correct,
            'game_mode': game_mode,
            'level': level
        }
        
        # Load existing answer logs
        try:
            with open('data/student_answers_log.json', 'r') as f:
                answers_data = json.load(f)
        except FileNotFoundError:
            answers_data = []
        
        # Add new answer log
        answers_data.append(answer_log)
        
        # Keep only last 500 answers to prevent file from growing too large
        if len(answers_data) > 500:
            answers_data = answers_data[-500:]
        
        # Save updated answer logs
        os.makedirs('data', exist_ok=True)
        with open('data/student_answers_log.json', 'w') as f:
            json.dump(answers_data, f, indent=2)
        
        # Emit real-time update to teachers (only if socketio is available)
        try:
            socketio.emit('student_answer', answer_log, room='teachers')
        except:
            pass  # SocketIO not available, skip real-time update
        
    except Exception as e:
        print(f"Error logging student answer: {e}")

def auto_save_progress():
    """Auto-save player progress if enabled in settings"""
    settings = get_current_game_settings()
    if not settings.get('auto_save', True):
        return
    
    # Only save if we have meaningful progress to save
    if not session.get('player_name') or not session.get('selected_level'):
        return
    
    try:
        progress_data = {
            'player_name': session.get('player_name', 'Anonymous'),
            'selected_level': session.get('selected_level'),
            'score': session.get('score', 0),
            'player_hp': session.get('player_hp', 100),
            'enemy_hp': session.get('enemy_hp', 50),
            'q_index': session.get('q_index', 0),
            'correct_answers': session.get('correct_answers', 0),
            'wrong_answers': session.get('wrong_answers', 0),
            'highest_unlocked': session.get('highest_unlocked', 1),
            'level_completed': session.get('level_completed', False),
            'character': session.get('character'),
            'timestamp': time.time()
        }
        
        # Save to auto-save file
        auto_save_file = f"data/autosave_{session.get('player_name', 'anonymous')}.json"
        os.makedirs('data', exist_ok=True)
        with open(auto_save_file, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, indent=2)
            
        print(f"Auto-saved progress for {session.get('player_name')}")
        
    except Exception as e:
        print(f"Auto-save failed: {e}")

def load_auto_save_progress(player_name):
    """Load auto-saved progress for a player"""
    try:
        auto_save_file = f"data/autosave_{player_name}.json"
        if not os.path.exists(auto_save_file):
            return None
            
        with open(auto_save_file, 'r', encoding='utf-8') as f:
            progress_data = json.load(f)
            
        # Check if save is recent (within 24 hours)
        if time.time() - progress_data.get('timestamp', 0) > 86400:
            return None
            
        return progress_data
        
    except Exception as e:
        print(f"Failed to load auto-save: {e}")
        return None

# Define the function to get questions for a specific level
def get_questions_for_level(level_number, levels):
    level_info = next((lvl for lvl in levels if lvl["level"] == level_number), None)
    if level_info:
        base_questions = [q for q in questions if q.get("id") in level_info["questions"]]
        
        # Filter by chapter if a specific chapter is selected
        if 'selected_chapter' in session:
            try:
                chapters_data = load_chapters()
                selected_chapter = next((ch for ch in chapters_data.get("chapters", []) 
                                       if ch.get("id") == session['selected_chapter']), None)
                if selected_chapter:
                    chapter_question_ids = set(selected_chapter.get("question_ids", []))
                    base_questions = [q for q in base_questions if q.get("id") in chapter_question_ids]
            except Exception as e:
                print(f"Error filtering by chapter: {e}")
        
        # Apply adaptive difficulty if enabled
        settings = get_current_game_settings()
        if settings.get('adaptive_difficulty', False):
            return apply_adaptive_difficulty(base_questions, level_number)
        return base_questions
    return []

def apply_adaptive_difficulty(base_questions, level_number):
    """Adjust question difficulty based on player performance"""
    # Calculate player's recent accuracy
    correct = session.get('correct_answers', 0)
    wrong = session.get('wrong_answers', 0)
    total = correct + wrong
    
    if total == 0:
        return base_questions  # No performance data yet
    
    accuracy = correct / total
    
    # Adjust difficulty based on performance
    if accuracy > 0.8:  # Player doing very well - increase difficulty
        # Try to get harder questions from higher levels
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                all_levels = json.load(f)
            
            harder_questions = []
            for lvl in all_levels:
                if lvl["level"] > level_number:
                    harder_questions.extend([q for q in questions if q.get("id") in lvl["questions"]])
            
            if harder_questions:
                # Mix 70% base questions with 30% harder questions
                base_count = int(len(base_questions) * 0.7)
                hard_count = min(len(harder_questions), len(base_questions) - base_count)
                return base_questions[:base_count] + random.sample(harder_questions, hard_count)
                
        except Exception:
            pass
            
    elif accuracy < 0.5:  # Player struggling - decrease difficulty
        # Try to get easier questions from lower levels
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                all_levels = json.load(f)
            
            easier_questions = []
            for lvl in all_levels:
                if lvl["level"] < level_number:
                    easier_questions.extend([q for q in questions if q.get("id") in lvl["questions"]])
            
            if easier_questions:
                # Mix 70% base questions with 30% easier questions
                base_count = int(len(base_questions) * 0.7)
                easy_count = min(len(easier_questions), len(base_questions) - base_count)
                return base_questions[:base_count] + random.sample(easier_questions, easy_count)
                
        except Exception:
            pass
    
    return base_questions  # No adjustment needed or possible


# Route for level selection
@app.route('/select_level', methods=['GET', 'POST'])
def select_level():
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except Exception:
        levels = []
    
    # Load chapters for level mode
    chapters_data = load_chapters()
    all_chapters = sorted(chapters_data.get("chapters", []), key=lambda x: x.get("order", 0))
    
    # Group levels by chapters
    chapter_levels = {}
    for chapter in all_chapters:
        chapter_id = chapter.get('id')
        level_range = chapter.get('level_range', [])
        chapter_levels[chapter_id] = {
            'chapter': chapter,
            'levels': [l for l in levels if l.get('level') in level_range]
        }

    # Determine highest unlocked level (simple: highest completed in session, or 1 if none)
    highest_unlocked = session.get('highest_unlocked', 1)
    # Optionally, you could use leaderboard or persistent storage for this

    if request.method == 'POST':
        selected_level = int(request.form.get('level', 1))
        selected_chapter = request.form.get('chapter_id')
        if selected_chapter:
            session['selected_chapter'] = int(selected_chapter)
        session['selected_level'] = selected_level
        # Reset session variables for a new game
        settings = get_current_game_settings()
        session['score'] = 0
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        # Set enemy HP to the current setting
        session['enemy_level'] = selected_level
        session['enemy_hp'] = settings['base_enemy_hp']
        
        # Ensure HP values are properly set and not zero
        if session['player_hp'] <= 0:
            session['player_hp'] = settings['base_player_hp']
        if session['enemy_hp'] <= 0:
            session['enemy_hp'] = settings['base_enemy_hp']
        
        # Log level selection
        log_analytics_event('level_selected', {
            'level': selected_level,
            'player_hp': settings['base_player_hp'],
            'enemy_hp': settings['base_enemy_hp']
        })
        

        session['q_index'] = 0
        session['feedback'] = None
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session["level_start_time"] = time.time()
        session["game_start_time"] = time.time()
        session['level_completed'] = False
        session['enemy_defeated'] = False  # Reset enemy defeated status for new level
        session['current_timer'] = settings.get('question_time_limit', 30)  # Use configurable time limit
        
        # Initialize enemy progression index based on selected level and current progress
        # Only reset to novice if playing level 1 or if no enemy index exists
        try:
            current_enemy_index = session.get('enemy_index')
            
            # If no enemy index exists, start with novice
            # Only reset to novice for level 1 if player hasn't progressed yet
            if current_enemy_index is None or (selected_level == 1 and current_enemy_index == 0):
                novice_idx = 0
                try:
                    with open('data/enemies.json', encoding='utf-8') as ef:
                        enemies_list = json.load(ef)
                    # Look for an enemy by name or level that indicates the novice gnome
                    for i, e in enumerate(enemies_list):
                        name = str(e.get('name', '')).strip().lower()
                        level = e.get('level')
                        if name == 'novice gnome' or level == 1:
                            novice_idx = i
                            break
                except Exception:
                    # If we can't read the file, default to index 0
                    novice_idx = 0
                session['enemy_index'] = int(novice_idx)
            else:
                # Set enemy to match the selected level
                try:
                    with open('data/enemies.json', encoding='utf-8') as ef:
                        enemies_list = json.load(ef)
                    
                    # Set enemy index to match the selected level
                    level_based_index = min(selected_level - 1, len(enemies_list) - 1)
                    session['enemy_index'] = level_based_index
                except Exception:
                    # Keep current enemy index if file can't be read
                    pass
        except Exception:
            session['enemy_index'] = 0
        # Auto-save initial game state
        auto_save_progress()
        
        # Check if student is logged in and has a character
        if session.get('is_student') and 'student_id' in session:
            students = load_students()
            student = next((s for s in students if s['id'] == session['student_id']), None)
            if student and 'selected_character' in student:
                # Student has a character, use it and go to game
                session['character'] = student['selected_character']
                return redirect(url_for('game'))
            else:
                # Student needs to choose a character
                session['redirect_after_character'] = 'game'
                return redirect(url_for('choose_character'))
        else:
            # Non-student user (guest) - always let them choose/change character for each new game
            session['redirect_after_character'] = 'game'
            return redirect(url_for('choose_character'))

    # Pass unlocked info to template (GET request) along with chapters
    return render_template('select_level.html', levels=levels, highest_unlocked=highest_unlocked, 
                         chapters=all_chapters, chapter_levels=chapter_levels)

# Route for the home page
@app.route('/')
def index():
    # Get settings to control feature visibility
    settings = get_current_game_settings()
    # Always redirect to level selection if no level is selected
    # If no selected level, show the merged landing page
    if 'selected_level' not in session:
        return render_template('index.html', settings=settings)
    # If a selected level exists, render the index page
    return render_template('index.html', settings=settings)


# Home route (explicit)
@app.route('/home')
def home():
    return render_template('home.html')


# How-to-play route (ensure it exists)
@app.route('/howto')
def howto():
    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)
    return render_template('howto.html', is_student=is_student)


@app.route('/choose_character', methods=['GET', 'POST'])
def choose_character():
    total_characters = 16
    
    # Check if student is logged in
    if session.get('is_student') and 'student_id' in session:
        student_id = session['student_id']
        students = load_students()
        student = next((s for s in students if s['id'] == student_id), None)
        
        if request.method == 'POST':
            char = request.form.get('character') or request.form.get('character_id')
            try:
                char_id = int(char)
            except (TypeError, ValueError):
                char_id = None
            if char_id and 1 <= char_id <= total_characters:
                # Save character to student profile
                if student:
                    student['selected_character'] = char_id
                    save_students(students)
                session['character'] = char_id
                flash(f'Character {char_id} selected!', 'success')
                # After choosing a character, redirect based on where they came from
                next_url = request.args.get('next') or session.get('redirect_after_character') or 'student_dashboard'
                session.pop('redirect_after_character', None)
                return redirect(url_for(next_url))
        
        # Check if student already has a character
        if student and 'selected_character' in student:
            session['character'] = student['selected_character']
            # Show current character selection
            return render_template('choose_character.html', 
                                 total_characters=total_characters,
                                 current_character=student['selected_character'],
                                 student_name=student['full_name'])
    
    # For non-logged in users or new character selection
    if request.method == 'POST':
        char = request.form.get('character') or request.form.get('character_id')
        try:
            char_id = int(char)
        except (TypeError, ValueError):
            char_id = None
        if char_id and 1 <= char_id <= total_characters:
            session['character'] = char_id
            # After choosing a character, start the game
            return redirect(url_for('game'))
    
    # Check if guest player has a current character to display
    current_char = session.get('character')
    return render_template('choose_character.html', 
                         total_characters=total_characters,
                         current_character=current_char)

# Route for the game page
@app.route('/game', methods=['GET', 'POST'])
def game():
    # Ensure player has selected a level and has basic session data initialized
    if 'selected_level' not in session or 'character' not in session:
        flash('Please select a level first.', 'warning')
        return redirect(url_for('select_level'))
    # Ensure player has a name
    if not session.get('player_name') or not session.get('character'):
        flash('Please set your name first.', 'warning')
        return redirect(url_for('index'))
    
    # Get settings (needed for both initialization and debug checks)
    settings = get_current_game_settings()
    # Initialize HP if not set (for new games)
    if 'player_hp' not in session or 'enemy_hp' not in session:
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        session['score'] = 0
        session['q_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['level_completed'] = False
        session['enemy_defeated'] = False
    
    # Check if the game is over (only after ensuring HP is initialized)
    if session.get('q_index', 0) >= len(questions):
        return redirect(url_for('result'))
    
    # Debug HP check
    current_hp = session.get('player_hp', 100)
    if settings.get('debug_mode', False):
        print(f"DEBUG: HP Check - current_hp: {current_hp}, q_index: {session.get('q_index', 0)}")
    
    if current_hp <= 0:
        if settings.get('debug_mode', False):
            print(f"DEBUG: Redirecting to you_lose due to HP <= 0 (HP: {current_hp})")
        return redirect(url_for('you_lose'))


    # Use selected level from session
    selected_level = session.get('selected_level', 1)
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except FileNotFoundError:
        return "Error: levels.json file not found."
    except json.JSONDecodeError as e:
        return f"Error: Failed to decode levels.json - {e}"

    # Only allow playing the selected level
    current_level = selected_level
    level_questions = get_questions_for_level(current_level, levels)
    print(f"DEBUG: Level {current_level} questions: {len(level_questions)} questions loaded")  # Debugging
    if not level_questions:
        flash(f'No questions available for level {current_level}. Please contact your teacher.', 'error')
        return redirect(url_for('select_level'))

    # Check if we've exceeded max questions
    questions_per_level = settings.get('questions_per_level', 10)
    
    # Mark enemy as defeated if HP <= 0 but continue playing
    if session.get('enemy_hp', BASE_ENEMY_HP) <= 0:
        session['enemy_defeated'] = True
        session['level_completed'] = True
    
    # Only end the level when all questions are answered
    if session['q_index'] >= questions_per_level:
        return redirect(url_for('result'))
    
    # Get question with proper bounds checking
    question_index = session['q_index'] % len(level_questions) if level_questions else 0
    if question_index >= len(level_questions):
        return redirect(url_for('result'))
    question = level_questions[question_index]

    # Always reload enemies from JSON for each game session
    try:
        with open('data/enemies.json', encoding='utf-8') as f:
            enemies = json.load(f)
    except Exception:
        enemies = []
    # Select enemy based on progression index, with fallbacks
    enemy = None
    enemy_index = session.get('enemy_index', 0)
    
    # Try to get enemy by progression index first
    if isinstance(enemies, list) and enemies and 0 <= int(enemy_index) < len(enemies):
        try:
            enemy = enemies[int(enemy_index)]
            print(f"DEBUG: Using enemy from index {enemy_index}: {enemy.get('name', 'Unknown')}")
        except Exception:
            enemy = None

    # Fallback 1: try to find an enemy that matches the current level
    if not enemy:
        enemy = next((e for e in enemies if e.get("level") == current_level), None)
        if enemy:
            print(f"DEBUG: Found level-based enemy for level {current_level}: {enemy.get('name', 'Unknown')}")
    
    # Fallback 2: use enemy based on level progression (level-1 as index)
    if not enemy and isinstance(enemies, list) and enemies:
        level_based_index = min(max(current_level - 1, 0), len(enemies) - 1)
        try:
            enemy = enemies[level_based_index]
            # Update session to match this enemy for consistency
            session['enemy_index'] = level_based_index
            print(f"DEBUG: Using level-based index {level_based_index} for level {current_level}: {enemy.get('name', 'Unknown')}")
        except Exception:
            enemy = None
    
    # Final fallback: default enemy
    if not enemy:
        enemy = {"name": "Unknown Enemy", "avatar": "❓", "taunt": "No enemies found for this level."}
    
    print(f"DEBUG: Final selected enemy for level {current_level}: {enemy.get('name', 'Unknown')} (enemy_index: {session.get('enemy_index')})")

    # Determine enemy image URL to mirror how player avatar images are used
    enemy_image = None
    # If the enemy JSON includes an explicit image field, use it (assume path under static/)
    if isinstance(enemy, dict) and enemy.get('image'):
        enemy_image = url_for('static', filename=enemy.get('image'))
    else:
        # Attempt to find an image in static/enemies matching the enemy name (safe fallback)
        # Create a safe filename from the enemy name (lowercase, replace spaces with underscores)
        if isinstance(enemy, dict) and enemy.get('name'):
            safe_base = ''.join(c if c.isalnum() else '_' for c in enemy.get('name').lower()).strip('_')
            # Check common extensions
            for ext in ['.png', '.jpg', '.jpeg', '.gif']:
                candidate = f'enemies/{safe_base}{ext}'
                candidate_path = os.path.join(os.path.dirname(__file__), 'static', candidate)
                if os.path.exists(candidate_path):
                    enemy_image = url_for('static', filename=candidate)
                    break

            # If not found by safe name, try fuzzy matching against files in static/enemies
            if not enemy_image:
                try:
                    enemy_tokens = [t for t in ''.join(c if c.isalnum() else ' ' for c in enemy.get('name').lower()).split() if t]
                    enemies_dir = os.path.join(os.path.dirname(__file__), 'static', 'enemies')
                    best_match = None
                    best_score = 0
                    if os.path.isdir(enemies_dir):
                        for fn in os.listdir(enemies_dir):
                            fn_lower = fn.lower()
                            name_part = os.path.splitext(fn_lower)[0]
                            file_tokens = [t for t in ''.join(c if c.isalnum() else ' ' for c in name_part).split() if t]
                            # score = number of matching tokens
                            score = sum(1 for t in enemy_tokens if t in file_tokens)
                            # if all tokens match, immediate perfect match
                            if score == len(enemy_tokens) and score > 0:
                                best_match = fn
                                best_score = score
                                break
                            if score > best_score:
                                best_score = score
                                best_match = fn
                        # accept best match if it matches at least one token
                        if best_match and best_score > 0:
                            candidate = f'enemies/{best_match}'
                            enemy_image = url_for('static', filename=candidate)
                except Exception:
                    enemy_image = None

    # Attach enemy_image into the template context

    # Initialize or reset the timer for the current question
    # Always reset timer on GET request
    if request.method == 'GET' or "level_start_time" not in session:
        session["level_start_time"] = time.time()
        # Initialize current_timer if not set
        if "current_timer" not in session:
            session["current_timer"] = settings.get('question_time_limit', 30)

    # Calculate remaining time
    settings = get_current_game_settings()
    elapsed = time.time() - session["level_start_time"]
    time_left = max(0, session.get("current_timer", settings.get('question_time_limit', 30)) - int(elapsed))

    # Debug timer info
    if settings.get('debug_mode', False):
        print(f"DEBUG: Timer - elapsed: {elapsed:.2f}s, time_left: {time_left}s, question_time_limit: {settings.get('question_time_limit', 30)}s, current_timer: {session.get('current_timer', 'not set')}")

    if time_left == 0:
        # Time expired → check timeout behavior setting
        settings = get_current_game_settings()
        session['wrong_answers'] = session.get('wrong_answers', 0) + 1  # Track timeouts as wrong answers
        
        # Check timeout behavior setting
        timeout_behavior = settings.get('timeout_behavior', 'penalty')
        
        if timeout_behavior == 'fail':
            # Immediate fail on timeout
            session['feedback'] = "⏳ Time's up! Timeout results in immediate failure."
            return redirect(url_for('you_lose'))
        else:
            # Apply penalty and continue (default behavior) - use base damage only, not multiplied by level
            timeout_damage = settings['base_damage']
            current_hp = session.get('player_hp', settings['base_player_hp'])
            
            # Ensure HP is properly initialized
            if 'player_hp' not in session:
                session['player_hp'] = settings['base_player_hp']
                current_hp = settings['base_player_hp']
            
            # Apply damage
            session['player_hp'] = current_hp - timeout_damage
            
            # Debug info
            if settings.get('debug_mode', False):
                print(f"DEBUG: Timeout - HP before: {current_hp}, damage: {timeout_damage}, HP after: {session['player_hp']}")
            
            # Check if player has failed due to low HP
            if session['player_hp'] <= 0:
                session['feedback'] = f"⏳ Time's up! You took {timeout_damage} damage and your HP reached 0. Game Over!"
                return redirect(url_for('you_lose'))
            
            session['feedback'] = f"⏳ Time's up! You took {timeout_damage} damage for running out of time. HP: {session['player_hp']}"
            session['q_index'] += 1
            session["level_start_time"] = time.time()  # Reset timer for the next question
            session["current_timer"] = max(10, session.get("current_timer", settings.get('question_time_limit', 30)) - 5)  # Deduct 5s for next question, min 10s
            
            # Auto-save progress after timeout
            auto_save_progress()
            
            return redirect(url_for('feedback'))

    if request.method == 'POST':
        user_answer = request.form.get('answer', '').strip().lower()
        correct_answer = question.get('answer', '').strip().lower()

        # Normalize keywords
        raw_keywords = question.get('keywords', [])
        if isinstance(raw_keywords, str):
            keywords = [k.strip().lower() for k in raw_keywords.split(',') if k.strip()]
        else:
            keywords = [str(k).strip().lower() for k in raw_keywords]

        # Calculate time taken to answer
        time_taken = time.time() - session["level_start_time"]

        # Determine damage and score based on time taken
        if time_taken <= 5:  # First 5 seconds → double damage and double score
            damage = 20
            score = 20
        elif time_taken <= 15:  # Within 15 seconds → regular damage and score
            damage = 10
            score = 10
        else:  # Remaining time → half damage and low score
            damage = 5
            score = 5

        # Only apply this block for the main game mode, not endless or test yourself
        if not session.get('endless_questions') and not session.get('test_questions'):
            # Get current game settings
            settings = get_current_game_settings()
            base_damage = settings['base_damage']
            points_correct = settings['points_correct']
            points_wrong = settings['points_wrong']
            
            # Debug logging if enabled
            if settings.get('debug_mode', False):
                print(f"[DEBUG] Player answer: '{user_answer}' for question: '{question.get('q', 'N/A')[:50]}...'")
                print(f"[DEBUG] Expected answer: '{correct_answer}', Keywords: {keywords}")
                print(f"[DEBUG] Time taken: {time_taken:.2f}s, Damage: {damage}, Score: {score}")
            
            # Get the question feedback
            question_feedback = question.get('feedback', 'No additional information available.')
            
            # Use fuzzy matching for answer checking
            is_correct, feedback_type, similarity_score = check_answer_fuzzy(user_answer, question)
            
            # If student and AI grading is enabled, use AI as fallback for uncertain answers
            if session.get('is_student') and session.get('ai_grading_enabled', False):
                # Use AI grading for short answers with low confidence (< 0.9)
                if question.get('type', 'short_answer') == 'short_answer' and not is_correct and similarity_score < 0.9:
                    try:
                        ai_result = grade_answer_with_ai(
                            question=question.get('q', ''),
                            correct_answer=correct_answer,
                            student_answer=user_answer,
                            confidence_threshold=75
                        )
                        if ai_result.get('correct', False) and ai_result.get('confidence', 0) >= 75:
                            is_correct = True
                            feedback_type = f"AI Grading: {ai_result.get('explanation', 'Accepted')}"
                            similarity_score = ai_result.get('confidence', 0) / 100.0
                    except Exception as e:
                        print(f"AI grading error: {e}")
            
            # Log student answer in real-time
            if 'student_id' in session:
                log_student_answer(
                    student_id=session['student_id'],
                    student_name=session.get('student_name', 'Unknown'),
                    question_id=question.get('id', 'unknown'),
                    question_text=question.get('q', ''),
                    student_answer=user_answer,
                    correct_answer=correct_answer,
                    is_correct=is_correct,
                    game_mode='adventure',
                    level=current_level
                )
            
            # Log answer attempt
            log_analytics_event('answer_submitted', {
                'question_id': question.get('id'),
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'time_taken': time_taken,
                'similarity_score': similarity_score,
                'level': current_level
            })
            
            if is_correct:
                # Calculate speed bonus if enabled
                speed_multiplier = 1.0
                speed_message = ""
                if settings.get('speed_bonus', True) and time_taken <= 10:
                    if time_taken <= 5:
                        speed_multiplier = 2.0
                        speed_message = " (🚀 Lightning bonus: x2 points!)"
                    elif time_taken <= 10:
                        speed_multiplier = 1.5
                        speed_message = " (⚡ Speed bonus: x1.5 points!)"
                
                points_awarded = int(points_correct * speed_multiplier)
                
                # Check if enemy is already defeated for bonus scoring
                enemy_already_defeated = session.get('enemy_defeated', False)
                
                if "Exact answer" in feedback_type or similarity_score == 1.0:
                    # Exact match gets triple damage or bonus points
                    if enemy_already_defeated:
                        # Enemy already defeated - give bonus points instead of damage
                        bonus_points = points_awarded * 2  # Double bonus for exact match after defeat
                        session['score'] += bonus_points
                        session['feedback'] = f"🏆 {feedback_type} Enemy already defeated! Bonus points: {bonus_points} in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                    else:
                        # Normal damage to enemy
                        session['score'] += points_awarded
                        session['enemy_hp'] -= (base_damage * 3)
                        # Ensure enemy HP doesn't go below 0
                        if session['enemy_hp'] <= 0:
                            session['enemy_hp'] = 0
                            session['enemy_defeated'] = True
                        session['feedback'] = f"🔥 {feedback_type} Triple damage: {base_damage*3} and {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                else:
                    # Fuzzy match gets regular damage or normal points
                    similarity_percent = int(similarity_score * 100)
                    if enemy_already_defeated:
                        # Enemy already defeated - give normal points (same as regular scoring)
                        session['score'] += points_awarded
                        session['feedback'] = f"🏆 {feedback_type} ({similarity_percent}% match) Enemy already defeated! Earned {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                    else:
                        # Normal damage to enemy
                        session['score'] += points_awarded
                        session['enemy_hp'] -= base_damage
                        # Ensure enemy HP doesn't go below 0
                        if session['enemy_hp'] <= 0:
                            session['enemy_hp'] = 0
                            session['enemy_defeated'] = True
                        session['feedback'] = f"🎯 {feedback_type} ({similarity_percent}% match) You dealt {base_damage} damage and earned {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                
                session["current_timer"] = min(settings.get('question_time_limit', 30) * 2, session.get("current_timer", settings.get('question_time_limit', 30)) + 5)  # Add 5s to timer for next question, max 2x base limit
                session['correct_answers'] = session.get('correct_answers', 0) + 1  # Track correct answers
            else:
                session['player_hp'] -= base_damage  # Use base damage only, not multiplied by level
                session['score'] -= points_wrong  # Deduct points for wrong answer
                session['feedback'] = f"❌ Incorrect! You took {base_damage} damage.<br><br>💡 {question_feedback}"
                session["current_timer"] = max(10, session.get("current_timer", settings.get('question_time_limit', 30)) - 5)  # Deduct 5s from timer for next question, min 10s
                session['wrong_answers'] = session.get('wrong_answers', 0) + 1  # Track wrong answers

        # Auto-save progress after each question
        auto_save_progress()

        # Reset timer for the next question
        session['level_start_time'] = time.time()

        # Move to the next question
        session['q_index'] += 1
        return redirect(url_for('feedback'))

    # Get current settings for display options
    settings = get_current_game_settings()
    
    # Generate dynamic taunt based on the current question
    enemy_taunt = generate_enemy_taunt(question, enemy.get('name', 'Unknown Enemy'))
    
    return render_template('game.html',
                           question=question,
                           score=session['score'],
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           enemy_defeated=session.get('enemy_defeated', False),
                           q_number=question_index + 1,
                           total=questions_per_level,  # Use configurable questions per level
                           level=current_level,
                           enemy=enemy,
                           enemy_image=enemy_image,
                           enemy_taunt=enemy_taunt,
                           time_left=time_left,
                           settings=settings)

# Route for the feedback page
@app.route('/feedback')
def feedback():
    # If player's HP reached 0, go straight to defeat
    if session.get('player_hp', 0) <= 0:
        return redirect(url_for('you_lose'))
    if not session.get('feedback'):
        return redirect(url_for('game'))  # Redirect to the game page if no feedback
    feedback = session['feedback']
    session['feedback'] = None  # Clear feedback after rendering
    return render_template('feedback.html',
                           feedback=feedback,
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           score=session['score'],
                           level=session['enemy_level'])

# Route for the result page
@app.route('/result')
def result():
    # Calculate the final score
    settings = get_current_game_settings()
    bonus = settings['level_bonus'] if session.get('level_completed', False) and session['player_hp'] > 0 else 0
    final_score = session['score'] + bonus

    # Save the player's performance to the leaderboard
    total_time = time.time() - session.get("game_start_time", time.time())
    save_leaderboard(
        player_name=session.get("player_name", "Anonymous"),
        score=final_score,
        total_time=total_time,
        correct_answers=session.get('correct_answers', 0),
        wrong_answers=session.get('wrong_answers', 0),
        game_mode="adventure",
        level=session.get('selected_level', 1)
    )
    
    # Save student progress if logged in
    if session.get('is_student') and session.get('student_id'):
        student_id = session.get('student_id')
        level = session.get('selected_level')
        correct_answers = session.get('correct_answers', 0)
        wrong_answers = session.get('wrong_answers', 0)
        total_questions = correct_answers + wrong_answers
        
        update_student_progress(
            student_id=student_id,
            game_type='adventure_mode',
            level=level,
            score=final_score,
            correct_answers=correct_answers,
            total_questions=total_questions,
            time_taken=total_time
        )

    # If the player died, show lose page
    if session.get('player_hp', 0) <= 0:
        return redirect(url_for('you_lose'))

    # Log level completion
    log_analytics_event('level_completed', {
        'level': session.get('selected_level', 1),
        'final_score': final_score,
        'player_hp': session.get('player_hp', 0),
        'enemy_hp': session.get('enemy_hp', 0),
        'total_time': total_time,
        'correct_answers': session.get('correct_answers', 0),
        'wrong_answers': session.get('wrong_answers', 0),
        'level_completed': session.get('level_completed', False)
    })
    
    # Check if player completed level 10 with 7+ correct answers - show credits
    current_level = session.get('selected_level', 1)
    correct_answers = session.get('correct_answers', 0)
    if current_level == 10 and correct_answers >= 7 and session.get('level_completed', False):
        return redirect(url_for('credits'))

    # If the level is completed, allow to select next level
    next_level = session.get('selected_level', 1) + 1
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except Exception:
        levels = []
    max_level = max([lvl['level'] for lvl in levels], default=1)
    
    # Calculate required accuracy based on settings
    settings = get_current_game_settings()
    total_questions = session.get('correct_answers', 0) + session.get('wrong_answers', 0)
    current_accuracy = (session.get('correct_answers', 0) / total_questions * 100) if total_questions > 0 else 0
    required_accuracy = settings.get('min_accuracy', 70)
    
    can_advance = (
        session.get('level_completed', False)
        and current_accuracy >= required_accuracy
        and next_level <= max_level
    )

    # Unlock the next level if completed AND meets accuracy requirement
    if session.get('level_completed', False):
        prev_highest = session.get('highest_unlocked', 1)
        
        # Only unlock next level if accuracy requirement is met
        if current_accuracy >= required_accuracy and next_level > prev_highest:
            session['highest_unlocked'] = next_level
        elif current_accuracy < required_accuracy:
            # Player completed level but didn't meet accuracy requirement
            # They can replay the same level but can't advance
            pass

        # Auto-save progress after level completion
        auto_save_progress()

        # Advance the enemy_index so a new enemy appears after each completed level
        try:
            with open('data/enemies.json', encoding='utf-8') as f:
                enemies_list = json.load(f)
        except Exception:
            enemies_list = []
        
        if isinstance(enemies_list, list) and enemies_list:
            current_idx = session.get('enemy_index', 0) or 0
            try:
                current_idx = int(current_idx)
            except Exception:
                current_idx = 0
            
            # Advance to next enemy, but don't exceed the list bounds
            next_idx = min(current_idx + 1, len(enemies_list) - 1)
            
            # Also ensure the enemy progression matches or exceeds the level progression
            # Each completed level should advance at least to that level's enemy
            level_based_idx = min(next_level - 1, len(enemies_list) - 1) if next_level <= max_level else len(enemies_list) - 1
            
            # Use the higher of the two indices to ensure proper progression
            final_idx = max(next_idx, level_based_idx)
            session['enemy_index'] = final_idx
            
            print(f"DEBUG: Enemy progression - Level: {next_level}, Current idx: {current_idx}, Next idx: {next_idx}, Level-based idx: {level_based_idx}, Final idx: {final_idx}")
        # If we've unlocked past the maximum level, the player beat the game
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                levels = json.load(f)
        except Exception:
            levels = []
        max_level = max([lvl['level'] for lvl in levels], default=1)
        if session.get('highest_unlocked', 1) > max_level:
            # Player has unlocked beyond max level -> they completed all levels
            return redirect(url_for('you_win'))

    return render_template('result.html', score=final_score,
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           outcome="Game Over",
                           level=session['enemy_level'],
                           enemy={"name": "Unknown Enemy", "avatar": "❓"},
                           can_advance=can_advance,
                           next_level=next_level)

# Route for the search functionality
from whoosh.qparser import QueryParser

@app.route('/search', methods=['GET'])
def search():
    query_text = request.args.get('q', '').strip()
    results = []
    
    if query_text:
        query_lower = query_text.lower()
        # Search through questions using multiple criteria
        for question in questions:
            match_found = False
            
            # Search in question text
            if query_lower in question.get('q', '').lower():
                match_found = True
            
            # Search in answer
            elif query_lower in question.get('answer', '').lower():
                match_found = True
            
            # Search in keywords
            elif question.get('keywords'):
                keywords = question['keywords']
                if isinstance(keywords, str):
                    if query_lower in keywords.lower():
                        match_found = True
                elif isinstance(keywords, list):
                    if any(query_lower in keyword.lower() for keyword in keywords):
                        match_found = True
            
            # Search in feedback/explanation
            elif question.get('feedback') and query_lower in question.get('feedback', '').lower():
                match_found = True
            
            if match_found:
                results.append(question)
    
    return render_template('search.html', results=results)

# Route for the leaderboard page
@app.route('/leaderboard')
def leaderboard():
    try:
        with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            all_data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        all_data = []

    # Separate leaderboards by game mode
    adventure_data = [entry for entry in all_data if entry.get("game_mode", "adventure") == "adventure"]
    test_yourself_data = [entry for entry in all_data if entry.get("game_mode") == "test_yourself"]
    endless_data = [entry for entry in all_data if entry.get("game_mode") == "endless"]
    
    # Get best score per player for each mode
    def get_best_scores(data, limit=50):
        player_best = {}
        for entry in data:
            player_name = entry.get("player", "Anonymous")
            if player_name not in player_best:
                player_best[player_name] = entry
            else:
                # Keep entry with higher score, or if same score, lower time
                current_best = player_best[player_name]
                if (entry["score"] > current_best["score"] or 
                    (entry["score"] == current_best["score"] and entry["time"] < current_best["time"])):
                    player_best[player_name] = entry
        # Sort by score (highest first), then by time (lowest first)
        sorted_players = sorted(player_best.values(), key=lambda x: (-x["score"], x["time"]))
        return sorted_players[:limit]
    
    # For adventure mode, create per-level leaderboards
    adventure_levels = {}
    for level_num in range(1, 11):  # Levels 1-10
        level_data = [entry for entry in adventure_data if entry.get("level") == level_num]
        adventure_levels[level_num] = get_best_scores(level_data, 10)  # Top 10 per level
    
    # Overall adventure leaderboard (all levels combined)
    adventure_leaderboard = get_best_scores(adventure_data, 50)
    
    test_yourself_leaderboard = get_best_scores(test_yourself_data, 50)
    endless_leaderboard = get_best_scores(endless_data, 50)

    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)

    return render_template("leaderboard.html", 
                         adventure_leaderboard=adventure_leaderboard,
                         adventure_levels=adventure_levels,
                         test_yourself_leaderboard=test_yourself_leaderboard,
                         endless_leaderboard=endless_leaderboard,
                         is_student=is_student)


@app.route('/guest_leaderboard')
def guest_leaderboard():
    try:
        with open(GUEST_LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            all_data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        all_data = []

    # Separate guest leaderboards by game mode
    adventure_data = [entry for entry in all_data if entry.get("game_mode", "adventure") == "adventure"]
    test_yourself_data = [entry for entry in all_data if entry.get("game_mode") == "test_yourself"]
    endless_data = [entry for entry in all_data if entry.get("game_mode") == "endless"]
    
    # Get best score per player for each mode
    def get_best_scores(data, limit=50):
        player_best = {}
        for entry in data:
            player_name = entry.get("player", "Anonymous")
            if player_name not in player_best:
                player_best[player_name] = entry
            else:
                # Keep entry with higher score, or if same score, lower time
                current_best = player_best[player_name]
                if (entry["score"] > current_best["score"] or 
                    (entry["score"] == current_best["score"] and entry["time"] < current_best["time"])):
                    player_best[player_name] = entry
        # Sort by score (highest first), then by time (lowest first)
        sorted_players = sorted(player_best.values(), key=lambda x: (-x["score"], x["time"]))
        return sorted_players[:limit]
    
    # For adventure mode, create per-level leaderboards
    adventure_levels = {}
    for level_num in range(1, 11):  # Levels 1-10
        level_data = [entry for entry in adventure_data if entry.get("level") == level_num]
        adventure_levels[level_num] = get_best_scores(level_data, 10)  # Top 10 per level
    
    # Overall adventure leaderboard (all levels combined)
    adventure_leaderboard = get_best_scores(adventure_data, 50)
    
    test_yourself_leaderboard = get_best_scores(test_yourself_data, 50)
    endless_leaderboard = get_best_scores(endless_data, 50)

    return render_template("guest_leaderboard.html", 
                         adventure_leaderboard=adventure_leaderboard,
                         adventure_levels=adventure_levels,
                         test_yourself_leaderboard=test_yourself_leaderboard,
                         endless_leaderboard=endless_leaderboard)


@app.route('/you_win')
def you_win():
    # Simple win page when all levels are completed
    # Reset enemy state to original when the player finishes the game
    try:
        settings = get_current_game_settings()
        session['enemy_index'] = 0
        session['enemy_hp'] = settings['base_enemy_hp']
        session.pop('enemy_level', None)
    except Exception:
        pass
    
    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)
    return render_template('you_win.html', is_student=is_student)


@app.route('/credits')
def credits():
    """Display credits page after completing level 10 with 7+ correct answers"""
    # Get final score and stats from session
    settings = get_current_game_settings()
    bonus = settings['level_bonus'] if session.get('level_completed', False) and session['player_hp'] > 0 else 0
    final_score = session.get('score', 0) + bonus
    correct_answers = session.get('correct_answers', 0)
    wrong_answers = session.get('wrong_answers', 0)
    total_time = time.time() - session.get("game_start_time", time.time())
    
    # Check if student is logged in
    is_student = session.get('is_student', False)
    
    return render_template('credits.html', 
                         final_score=final_score,
                         correct_answers=correct_answers,
                         wrong_answers=wrong_answers,
                         total_time=round(total_time, 2),
                         is_student=is_student)


@app.route('/you_lose')
def you_lose():
    # Reset only the current level progress, not the entire game
    try:
        settings = get_current_game_settings()
        current_level = session.get('selected_level', 1)
        
        # Reset level-specific progress but keep level selection
        session['q_index'] = 0
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        session['enemy_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['score'] = 0
        session['level_completed'] = False
        session['enemy_defeated'] = False
        session.pop('feedback', None)
        
        # Keep the selected level so they can retry the same level
        session['selected_level'] = current_level
        session['enemy_level'] = current_level
        
        # Log the level failure for analytics
        log_analytics_event('level_failed', {
            'level': current_level,
            'reason': 'hp_zero'
        })
        
    except Exception as e:
        print(f"Error in you_lose route: {e}")
    
    return render_template('you_lose.html')

@app.route('/quit_game')
def quit_game():
    """Handle player quitting mid-game"""
    try:
        current_level = session.get('selected_level', 1)
        current_score = session.get('score', 0)
        questions_answered = session.get('q_index', 0)
        
        # Log the quit action for analytics
        log_analytics_event('game_quit', {
            'level': current_level,
            'score': current_score,
            'questions_answered': questions_answered,
            'reason': 'player_quit'
        })
        
        # Clear current game progress but keep player info
        game_progress_keys = ['q_index', 'player_hp', 'enemy_hp', 'score', 
                             'correct_answers', 'wrong_answers', 'level_completed', 
                             'enemy_defeated', 'feedback', 'current_timer']
        
        for key in game_progress_keys:
            session.pop(key, None)
        
        flash(f'Game quit. You scored {current_score} points on Level {current_level}.', 'info')
        
    except Exception as e:
        print(f"Error in quit_game route: {e}")
        flash('Game ended.', 'info')
    
    return redirect(url_for('select_level'))

@app.route('/quit_test_yourself')
def quit_test_yourself():
    """Handle player quitting Test Yourself mode - auto-save to leaderboard"""
    try:
        questions_answered = session.get('test_q_index', 0)
        correct_answers = session.get('test_correct', 0)
        total_questions = len(session.get('test_question_ids', []))
        test_start_time = session.get('test_start_time', time.time())
        time_taken = time.time() - test_start_time
        score = correct_answers * 10  # Same scoring as normal completion
        player_name = session.get('player_name', 'Anonymous')
        
        # Auto-save to leaderboard when quitting
        save_leaderboard(
            player_name=player_name,
            score=score,
            total_time=time_taken,
            correct_answers=correct_answers,
            wrong_answers=questions_answered - correct_answers,
            game_mode="test_yourself"
        )
        
        # Save student progress if logged in
        if session.get('is_student') and session.get('student_id'):
            student_id = session.get('student_id')
            update_student_progress(
                student_id=student_id,
                game_type='test_yourself',
                level=None,
                score=score,
                correct_answers=correct_answers,
                total_questions=total,
                time_taken=time_taken
            )
        
        # Log the quit action for analytics
        log_analytics_event('test_yourself_quit', {
            'questions_answered': questions_answered,
            'correct_answers': correct_answers,
            'score': score,
            'reason': 'player_quit'
        })
        
        # Clear all test yourself session data completely
        reset_test_yourself_session()
        
        accuracy = (correct_answers / questions_answered * 100) if questions_answered > 0 else 0
        flash(f'Test Yourself completed (quit). Score: {score}, Accuracy: {accuracy:.1f}% - Saved to leaderboard!', 'success')
        
    except Exception as e:
        print(f"Error in quit_test_yourself route: {e}")
        flash('Test session ended.', 'info')
    
    return redirect(url_for('leaderboard'))

@app.route('/quit_endless')
def quit_endless():
    """Handle player quitting Endless mode - auto-save to leaderboard"""
    try:
        total_answered = session.get('endless_total_answered', 0)
        correct_answers = session.get('endless_correct', 0)
        wrong_answers = session.get('endless_wrong', 0)
        final_score = session.get('endless_score', 0)
        highest_streak = session.get('endless_highest_streak', 0)
        endless_start_time = session.get('endless_start_time', time.time())
        total_time = time.time() - endless_start_time
        player_name = session.get('player_name', 'Anonymous')
        
        # Auto-save to leaderboard when quitting
        save_leaderboard(
            player_name=player_name,
            score=final_score,
            total_time=total_time,
            correct_answers=correct_answers,
            wrong_answers=wrong_answers,
            game_mode="endless"
        )
        
        # Save student progress if logged in
        if session.get('is_student') and session.get('student_id'):
            student_id = session.get('student_id')
            update_student_progress(
                student_id=student_id,
                game_type='endless_mode',
                level=None,
                score=final_score,
                correct_answers=correct_answers,
                total_questions=total_answered,
                time_taken=total_time
            )
        
        # Log the quit action for analytics
        log_analytics_event('endless_mode_quit', {
            'questions_answered': total_answered,
            'final_score': final_score,
            'max_streak': highest_streak,
            'reason': 'player_quit'
        })
        
        # Clear all endless mode session data completely
        reset_endless_mode_session()
        
        flash(f'Endless Mode completed (quit). Score: {final_score}, {total_answered} questions - Saved to leaderboard!', 'success')
        
    except Exception as e:
        print(f"Error in quit_endless route: {e}")
        flash('Endless session ended.', 'info')
    
    return redirect(url_for('leaderboard'))

# ------------------- TEST YOURSELF MODE -------------------
import random

@app.route('/test_yourself', methods=['GET', 'POST'])
def test_yourself():
    # Check if Test Yourself mode is enabled
    settings = get_current_game_settings()
    if not settings.get('test_yourself_enabled', True):
        flash('Test Yourself mode is currently disabled.', 'error')
        return redirect(url_for('index'))
    
    # Reset test state for a true new start (GET with ?new=1) or if no session data exists
    if (request.method == 'GET' and request.args.get('new') == '1') or not session.get('test_question_ids'):
        # Completely reset session to ensure clean start
        reset_test_yourself_session()
        session['test_user_answers'] = []
        # ...existing code...
        
        # Use questions from test_yourself pool
        test_pool_questions = get_questions_for_pool('test_yourself')
        if not test_pool_questions:
            test_pool_questions = questions  # Fallback to all questions
        
        valid_questions = [q for q in test_pool_questions if q.get('q') and str(q.get('q')).strip()]
        if not valid_questions:
            session['test_question_ids'] = []
        elif len(valid_questions) >= 40:
            # Use random.sample to guarantee no duplicates (returns unique selection)
            selected = random.sample(valid_questions, 40)
            # Extract IDs and convert to set for guaranteed uniqueness
            question_ids_set = set(q['id'] for q in selected)
            # Convert back to list and shuffle
            question_ids_list = list(question_ids_set)
            random.shuffle(question_ids_list)
            session['test_question_ids'] = question_ids_list
        else:
            # If fewer than 40 questions, use all available without repeats
            # Use set to ensure absolute uniqueness even with small pools
            unique_questions = list({q['id']: q for q in valid_questions}.values())
            random.shuffle(unique_questions)
            session['test_question_ids'] = [q['id'] for q in unique_questions]
        
        # Debug: Verify uniqueness using set comparison
        question_ids = session['test_question_ids']
        unique_ids = set(question_ids)
        print(f"[DEBUG TEST INIT] Selected {len(question_ids)} questions, {len(unique_ids)} unique IDs (set-verified)")
        if len(question_ids) != len(unique_ids):
            duplicate_ids = [id for id in unique_ids if question_ids.count(id) > 1]
            print(f"[CRITICAL TEST INIT] Duplicate question IDs found despite set conversion: {duplicate_ids}")
            # Force fix by using only unique IDs
            session['test_question_ids'] = list(unique_ids)
            random.shuffle(session['test_question_ids'])
        
        session['test_q_index'] = 0
        session['test_correct'] = 0
        session['test_start_time'] = time.time()
        session['test_time_limit'] = 60 * 60  # 1 hour in seconds

    # Calculate timer
    total_seconds_left = max(0, int(session.get('test_time_limit', 3600) - (time.time() - session.get('test_start_time', time.time()))))
    time_left_min = total_seconds_left // 60
    time_left_sec = total_seconds_left % 60
    q_index = session.get('test_q_index', 0)
    test_question_ids = session.get('test_question_ids', [])
    
    # Debug: Print current state
    print(f"[DEBUG TEST] GET request - q_index={q_index}, total_test_questions={len(test_question_ids)}")
    if q_index < len(test_question_ids):
        print(f"[DEBUG TEST] Current question ID: {test_question_ids[q_index]}")
    
    # Rebuild the test_questions list from global questions using IDs
    if not questions:
        flash('No questions available. Please contact your teacher.', 'error')
        return redirect(url_for('index'))
    
    try:
        id_to_question = {q['id']: q for q in questions}
        test_questions = [id_to_question[qid] for qid in test_question_ids if qid in id_to_question]
        
        # Store only question count, not full IDs list to reduce session size
        if 'test_total_questions' not in session:
            session['test_total_questions'] = len(test_question_ids)
    except Exception as e:
        print(f"[ERROR] Failed to rebuild test_questions: {e}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))
    
    # Check if we've reached the end or time is up (40 questions = indices 0-39)
    # Using > 39 instead of >= 40 to ensure question 40 (index 39) is displayed
    if not test_questions or q_index > 39 or total_seconds_left <= 0:
        print(f"[DEBUG] REDIRECT TO RESULT: test_q_index={q_index}, test_questions={len(test_questions)}, total_seconds_left={total_seconds_left}, test_user_answers={len(session.get('test_user_answers', []))}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))

    # Skip invalid questions
    while q_index < len(test_questions) and q_index < 40:
        try:
            if test_questions[q_index].get('q') and str(test_questions[q_index].get('q')).strip():
                break
            q_index += 1
            session['test_q_index'] = q_index
        except (IndexError, KeyError):
            q_index += 1
            session['test_q_index'] = q_index
    
    # Final check after skipping invalid questions
    if q_index > 39:
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))
    
    # Safety check for question existence
    try:
        question = test_questions[q_index]
        if not question or not question.get('q') or not str(question.get('q')).strip():
            raise IndexError("Invalid question")
    except (IndexError, KeyError) as e:
        print(f"[ERROR] Question access error at index {q_index}: {e}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))

    correct_count = session.get('test_correct', 0)
    if request.method == 'POST':
        try:
            user_answer = request.form.get('answer', '').strip().lower()
            correct_answer = question.get('answer', '').strip().lower()
            # Normalize keywords
            raw_keywords = question.get('keywords', [])
            if isinstance(raw_keywords, str):
                keywords = [k.strip().lower() for k in raw_keywords.split(',') if k.strip()]
            else:
                keywords = [str(k).strip().lower() for k in raw_keywords]
            # Use fuzzy matching for test mode
            is_correct, feedback_type, similarity_score = check_answer_fuzzy(user_answer, question)
            
            # If student and AI grading is enabled, use AI as fallback for uncertain answers
            if session.get('is_student') and session.get('ai_grading_enabled', False):
                # Use AI grading for short answers with low confidence (< 0.9)
                if question.get('type', 'short_answer') == 'short_answer' and not is_correct and similarity_score < 0.9:
                    try:
                        ai_result = grade_answer_with_ai(
                            question=question.get('q', ''),
                            correct_answer=correct_answer,
                            student_answer=user_answer,
                            confidence_threshold=75
                        )
                        if ai_result.get('correct', False) and ai_result.get('confidence', 0) >= 75:
                            is_correct = True
                            feedback_type = f"AI Grading: {ai_result.get('explanation', 'Accepted')}"
                            similarity_score = ai_result.get('confidence', 0) / 100.0
                    except Exception as e:
                        print(f"AI grading error: {e}")
            
            # Log student answer in real-time
            if 'student_id' in session:
                log_student_answer(
                    student_id=session['student_id'],
                    student_name=session.get('student_name', 'Unknown'),
                    question_id=question.get('id', 'unknown'),
                    question_text=question.get('q', ''),
                    student_answer=user_answer,
                    correct_answer=correct_answer,
                    is_correct=is_correct,
                    game_mode='test_yourself'
                )
            
            session['test_user_answers'].append({
                'question': question.get('q', '')[:150],  # Further truncate questions
                'user_answer': user_answer[:100],  # Truncate user answers
                'correct_answer': correct_answer[:100],  # Truncate correct answers
                'correct': is_correct,
                'feedback': question.get('feedback', '')[:150],  # Reduce feedback size
                'match_type': feedback_type[:50] if isinstance(feedback_type, str) else str(feedback_type)[:50],
                'similarity': round(similarity_score, 2) if similarity_score else 0
            })
            # Keep only essential answers, limit to 40
            if len(session['test_user_answers']) > 40:
                session['test_user_answers'] = session['test_user_answers'][-40:]
            if is_correct:
                session['test_correct'] = correct_count + 1
            
            # Increment question index
            new_q_index = q_index + 1
            session['test_q_index'] = new_q_index
            
            # Debug logging
            print(f"[DEBUG TEST POST] Answered Q{q_index + 1} (ID={question.get('id')}), correct={is_correct}")
            print(f"[DEBUG TEST POST] Moving to next: new_q_index={new_q_index}, total_questions={len(test_question_ids)}")
            if new_q_index < len(test_question_ids):
                print(f"[DEBUG TEST POST] Next question ID will be: {test_question_ids[new_q_index]}")
            
            return redirect(url_for('test_yourself'))
        except Exception as e:
            print(f"[ERROR] Test yourself mode POST error: {str(e)}")
            # Force game over on error to prevent crash
            session['test_q_index'] = 40
            return redirect(url_for('test_yourself_result'))

    return render_template('test_yourself.html',
                          question=question,
                          q_number=q_index + 1,
                          q_index=q_index,
                          test_questions=test_questions,
                          correct_count=correct_count,
                          time_left_min=time_left_min,
                          time_left_sec=time_left_sec,
                          total_seconds_left=total_seconds_left)

@app.route('/test_yourself_result')
def test_yourself_result():
    # Check if Test Yourself mode is enabled
    settings = get_current_game_settings()
    if not settings.get('test_yourself_enabled', True):
        flash('Test Yourself mode is currently disabled.', 'error')
        return redirect(url_for('index'))
    
    # Use stored total or question IDs length as fallback
    total = session.get('test_total_questions', len(session.get('test_question_ids', [])))
    correct = session.get('test_correct', 0)
    percent = int((correct / total) * 100) if total else 0
    passed = percent >= 75
    # Get user answers for review
    user_answers = session.get('test_user_answers', [])
    
    # Calculate time and score
    test_start_time = session.get('test_start_time', time.time())
    time_taken = time.time() - test_start_time
    score = correct * 10  # Simple scoring system
    
    # Get player name (student name or guest name)
    if session.get('is_student'):
        player_name = session.get('student_name', 'Anonymous')
    else:
        player_name = session.get('player_name', 'Anonymous')
    
    # Save to leaderboard for ALL players (students and guests)
    save_leaderboard(
        player_name=player_name,
        score=score,
        total_time=time_taken,
        correct_answers=correct,
        wrong_answers=total - correct,
        game_mode="test_yourself"
    )
    
    # Save student progress if logged in as student
    if session.get('is_student') and session.get('student_id'):
        student_id = session.get('student_id')
        update_student_progress(
            student_id=student_id,
            game_type='test_yourself',
            level=None,
            score=score,
            correct_answers=correct,
            total_questions=total,
            time_taken=time_taken
        )
    
    # Clear session state for test mode
    session.pop('test_question_ids', None)
    session.pop('test_q_index', None)
    session.pop('test_correct', None)
    session.pop('test_start_time', None)
    session.pop('test_time_limit', None)
    session.pop('test_user_answers', None)
    return render_template('test_yourself_result.html',
                          correct_count=correct,
                          percent=percent,
                          passed=passed,
                          user_answers=user_answers)

# ------------------- ENDLESS MODE -------------------

# Load game settings at startup
def load_initial_settings():
    """Load initial game settings or use defaults"""
    try:
        settings_file = 'data/game_settings.json'
        with open(settings_file, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        return settings
    except (FileNotFoundError, json.JSONDecodeError):
        # Create default settings file if it doesn't exist
        default_settings = {
            'base_player_hp': 100,
            'base_enemy_hp': 50,
            'base_damage': 10,
            'question_time_limit': 30,
            'questions_per_level': 10
        }
        try:
            os.makedirs('data', exist_ok=True)
            with open('data/game_settings.json', 'w', encoding='utf-8') as f:
                json.dump(default_settings, f, indent=2)
        except Exception:
            pass
        return default_settings

# Initialize settings
initial_settings = load_initial_settings()
BASE_DAMAGE = initial_settings.get('base_damage', 10)
BASE_ENEMY_HP = initial_settings.get('base_enemy_hp', 50)
LEVEL_TIME_LIMIT = initial_settings.get('question_time_limit', 30)

# Function to generate dynamic enemy taunts based on question
def generate_enemy_taunt(question, enemy_name):
    """Generate a dynamic taunt based on the question content"""
    question_text = question.get('q', '').lower()
    keywords = question.get('keywords', [])
    
    # Extract key topics from question
    if isinstance(keywords, str):
        keywords = [k.strip().lower() for k in keywords.split(',')]

    taunt_templates = {
        'ls': [
            "Can you even list a directory?",
            "Let's see if you know what 'ls' does!",
            "Basic commands? This should be easy... or not!"
        ],
        'cd': [
            "Lost in the filesystem already?",
            "Can you navigate directories?",
            "Where do you think you're going?"
        ],
        'chmod': [
            "Permissions confuse you, don't they?",
            "Can you handle file permissions?",
            "Let's test your permission knowledge!"
        ],
        'systemctl': [
            "Service management is my domain!",
            "Can you control system services?",
            "Let's see your systemctl skills!"
        ],
        'firewall': [
            "Your firewall knowledge is weak!",
            "Can you protect this system?",
            "Let's test your security skills!"
        ],
        'user': [
            "User management is tricky, isn't it?",
            "Can you handle users and groups?",
            "Let's see if you can manage users!"
        ],
        'network': [
            "Networking will be your downfall!",
            "Can you configure network settings?",
            "Let's test your network knowledge!"
        ],
        'mount': [
            "Can you mount filesystems?",
            "Storage management is complex!",
            "Let's see your mounting skills!"
        ],
        'selinux': [
            "SELinux is too advanced for you!",
            "Can you handle security contexts?",
            "Security-Enhanced Linux will defeat you!"
        ],
        'lvm': [
            "Logical volumes will confuse you!",
            "Can you manage LVM?",
            "Storage management is my specialty!"
        ],
        'cron': [
            "Can you schedule tasks?",
            "Time-based jobs are tricky!",
            "Let's test your automation skills!"
        ],
        'package': [
            "Package management is complex!",
            "Can you install software?",
            "Let's see your package skills!"
        ],
        'grep': [
            "Can you search through text?",
            "Let's test your pattern matching!",
            "Grep will be your challenge!"
        ],
        'find': [
            "Can you find files?",
            "Let's see your search skills!",
            "File searching will defeat you!"
        ],
        'tar': [
            "Archiving is too complex for you!",
            "Can you handle tar archives?",
            "Let's test your compression knowledge!"
        ],
        'dns': [
            "DNS will confuse you!",
            "Can you resolve hostnames?",
            "Let's test your name resolution skills!"
        ],
        'boot': [
            "Boot processes are tricky!",
            "Can you fix boot issues?",
            "System startup will challenge you!"
        ]
    }
    
    # Find matching topic
    for keyword in keywords:
        for topic, taunts in taunt_templates.items():
            if topic in keyword or topic in question_text:
                import random
                return random.choice(taunts)
    
    # Check question text for topics
    for topic, taunts in taunt_templates.items():
        if topic in question_text:
            import random
            return random.choice(taunts)
    
    # Generic taunts based on enemy name
    generic_taunts = [
        f"{enemy_name} challenges your knowledge!",
        "This question will test your skills!",
        "Can you answer this correctly?",
        "Let's see what you know!",
        "Prove your expertise!"
    ]
    
    import random
    return random.choice(generic_taunts)

# Define the schema for the Whoosh search index
schema = Schema(
    id=ID(stored=True, unique=True),
    question=TEXT(stored=True),
    answer=TEXT(stored=True),
    keywords=TEXT(stored=True)
)

# Create or open the Whoosh index directory
def create_or_open_index():
    """Open the Whoosh index if valid; otherwise recreate it."""
    try:
        if not os.path.exists("indexdir"):
            os.mkdir("indexdir")
            return index.create_in("indexdir", schema)
        try:
            return index.open_dir("indexdir")
        except Exception as e:
            print(f"Whoosh index open failed: {e}. Recreating indexdir...")
            try:
                # Remove the corrupted indexdir and recreate
                import shutil
                shutil.rmtree("indexdir")
            except Exception:
                pass
            os.mkdir("indexdir")
            return index.create_in("indexdir", schema)
    except Exception as e:
        print(f"Unexpected error creating/opening indexdir: {e}")
        raise

ix = create_or_open_index()

# Initialize the Flask app
app = Flask(__name__)
import copy

app.secret_key = "unix_rpg_secret"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Initialize Flask-SocketIO
try:
    socketio = SocketIO(app, cors_allowed_origins="*")
except Exception as e:
    print(f"Warning: SocketIO initialization failed: {e}")
    socketio = None

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.before_request
def check_session_timeout():
    """Check if session has timed out based on settings"""
    settings = get_current_game_settings()
    session_timeout_minutes = settings.get('session_timeout', 30)
    
    # Skip timeout check for certain routes
    excluded_routes = ['static', 'teacher_login', 'teacher_logout']
    if request.endpoint in excluded_routes:
        return
    
    # Check if user has activity timestamp
    if 'last_activity' in session:
        last_activity = session['last_activity']
        current_time = time.time()
        timeout_seconds = session_timeout_minutes * 60
        
        if current_time - last_activity > timeout_seconds:
            # Session timed out - clear session and redirect
            session.clear()
            flash(f'Your session timed out after {session_timeout_minutes} minutes of inactivity.', 'warning')
            return redirect(url_for('index'))
    
    # Update last activity timestamp
    session['last_activity'] = time.time()

# Teacher authentication decorator
def teacher_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('teacher_logged_in'):
            return redirect(url_for('teacher_login'))
        return f(*args, **kwargs)
    return decorated_function

# Student authentication decorator
def student_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('is_student'):
            return redirect(url_for('student_login'))
        return f(*args, **kwargs)
    return decorated_function

# Helper function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# AI Integration Functions
def extract_json_from_response(response):
    """Extract JSON from AI response, handling markdown code blocks and truncation"""
    import re
    
    # If the response contains markdown code blocks, extract the JSON
    json_pattern = r'```(?:json)?\s*(\[.*\]|\{.*\})\s*```'
    match = re.search(json_pattern, response, re.DOTALL)
    
    if match:
        return match.group(1).strip()
    
    # For JSON arrays, find the opening [ and matching closing ]
    start_idx = response.find('[')
    if start_idx != -1:
        bracket_count = 0
        for i, char in enumerate(response[start_idx:], start_idx):
            if char == '[':
                bracket_count += 1
            elif char == ']':
                bracket_count -= 1
                if bracket_count == 0:
                    return response[start_idx:i+1].strip()
    
    # For JSON objects, find the opening { and matching closing }
    start_idx = response.find('{')
    if start_idx != -1:
        brace_count = 0
        for i, char in enumerate(response[start_idx:], start_idx):
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    return response[start_idx:i+1].strip()
    
    # If no JSON structure found, return the response as-is
    return response.strip()

def fix_incomplete_json(json_str):
    """Attempt to fix incomplete JSON by adding missing closing brackets/braces"""
    json_str = json_str.strip()
    
    # Count opening and closing brackets/braces
    open_brackets = json_str.count('[')
    close_brackets = json_str.count(']')
    open_braces = json_str.count('{')
    close_braces = json_str.count('}')
    
    # Add missing closing brackets
    while close_brackets < open_brackets:
        json_str += ']'
        close_brackets += 1
    
    # Add missing closing braces
    while close_braces < open_braces:
        json_str += '}'
        close_braces += 1
    
    # If the JSON ends with a comma, try to remove it
    if json_str.rstrip().endswith(','):
        json_str = json_str.rstrip()[:-1]
    
    return json_str

def call_ai_api(prompt, max_tokens=1000):
    """Call AI API (OpenAI or Gemini) with error handling"""
    if AI_PROVIDER == "gemini":
        return call_gemini_api(prompt, max_tokens)
    elif AI_PROVIDER == "openai":
        return call_openai_api(prompt, max_tokens)
    else:
        return {"error": "Invalid AI provider configured"}

def is_ai_configured():
    """Check if AI API key is properly configured"""
    if AI_PROVIDER == "gemini":
        return bool(GEMINI_API_KEY and GEMINI_API_KEY.strip() and GEMINI_API_KEY != 'your-gemini-api-key-here' and len(GEMINI_API_KEY) > 10)
    elif AI_PROVIDER == "openai":
        return bool(OPENAI_API_KEY and OPENAI_API_KEY.strip() and OPENAI_API_KEY != 'your-openai-api-key-here' and len(OPENAI_API_KEY) > 10)
    return False

def get_ai_config_error_message():
    """Get user-friendly error message for missing AI configuration"""
    if AI_PROVIDER == "gemini":
        return "Gemini API key not configured. Please set GEMINI_API_KEY in config.py. Get a free key at https://makersuite.google.com/app/apikey"
    elif AI_PROVIDER == "openai":
        return "OpenAI API key not configured. Please set OPENAI_API_KEY in config.py. Get a key at https://platform.openai.com/api-keys"
    return "AI provider not configured properly in config.py"

def call_openai_api(prompt, max_tokens=1000):
    """Call OpenAI API with error handling"""
    try:
        headers = {
            'Authorization': f'Bearer {OPENAI_API_KEY}',
            'Content-Type': 'application/json',
        }
        data = {
            'model': OPENAI_MODEL,
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': max_tokens,
            'temperature': 0.7
        }
        response = requests.post('https://api.openai.com/v1/chat/completions', 
                               headers=headers, json=data, timeout=60)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error calling OpenAI API: {str(e)}"

def call_gemini_api(prompt, max_tokens=1000):
    """Call Google Gemini API with error handling"""
    try:
        if not GEMINI_API_KEY or GEMINI_API_KEY == 'your-gemini-api-key-here':
            return "Gemini API Error: API key not configured. Please set GEMINI_API_KEY in config.py"
        
        headers = {
            'Content-Type': 'application/json'
        }
        
        # Gemini API endpoint
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}'
        
        data = {
            'contents': [{
                'parts': [{'text': prompt}]
            }],
            'generationConfig': {
                'maxOutputTokens': max_tokens,
                'temperature': 0.7
            }
        }
        
        print(f"DEBUG: Calling Gemini API with model: {GEMINI_MODEL}")
        print(f"DEBUG: API Key configured: {bool(GEMINI_API_KEY and len(GEMINI_API_KEY) > 20)}")
        print(f"DEBUG: Max tokens requested: {max_tokens}")
        
        response = requests.post(url, headers=headers, json=data, timeout=60)
        
        print(f"DEBUG: Gemini API response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"DEBUG: Full Gemini response structure: {json.dumps(result, indent=2)[:500]}")
            
            if 'candidates' in result and len(result['candidates']) > 0:
                candidate = result['candidates'][0]
                
                # Check if response was blocked
                if 'finishReason' in candidate and candidate['finishReason'] != 'STOP':
                    return f"Gemini API Error: Response blocked or incomplete (reason: {candidate.get('finishReason', 'unknown')})"
                
                if 'content' in candidate and 'parts' in candidate['content']:
                    response_text = candidate['content']['parts'][0]['text']
                    print(f"DEBUG: Response length: {len(response_text)} characters")
                    print(f"DEBUG: Response preview: {response_text[:200]}")
                    return response_text
                else:
                    print(f"DEBUG: Unexpected candidate structure: {candidate}")
                    return "Gemini API Error: Response missing content or parts"
            else:
                print(f"DEBUG: No candidates in response: {result}")
                return f"Gemini API Error: No response generated. Full response: {json.dumps(result)[:300]}"
        else:
            error_details = response.text
            print(f"DEBUG: Gemini API Error Details: {error_details}")
            
            if response.status_code == 400:
                return "Gemini API Error: Invalid API key or request format. Please check your API key and try again."
            elif response.status_code == 403:
                return "Gemini API Error: API key doesn't have permission or quota exceeded."
            elif response.status_code == 429:
                return "Gemini API Error: Rate limit exceeded. Please try again later."
            else:
                return f"Gemini API Error: {response.status_code} - {error_details}"
            
    except Exception as e:
        print(f"DEBUG: Exception in call_gemini_api: {str(e)}")
        return f"Error calling Gemini API: {str(e)}"

def extract_text_from_file(file_path):
    """Extract text content from uploaded files"""
    try:
        if file_path.endswith('.txt') or file_path.endswith('.md'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.endswith('.pdf'):
            # You would need to install PyPDF2: pip install PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text
            except ImportError:
                return "PDF support requires PyPDF2. Install with: pip install PyPDF2"
        elif file_path.endswith('.docx'):
            # You would need to install python-docx: pip install python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                return "DOCX support requires python-docx. Install with: pip install python-docx"
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading file: {str(e)}"

def generate_questions_with_ai(content, topic, difficulty, question_count, context="", question_types=None):
    """Generate questions using AI based on content"""
    
    # Check if AI is configured
    if not is_ai_configured():
        return {"error": get_ai_config_error_message()}
    
    # Default question types if none specified
    if question_types is None:
        question_types = ["short_answer", "multiple_choice", "true_false"]
    
    # Calculate appropriate content length based on question count
    # More questions need more tokens for response, so reduce content accordingly
    base_content_limit = 6000
    content_limit = max(2000, base_content_limit - (question_count * 200))
    
    # Truncate content if needed
    if len(content) > content_limit:
        content = content[:content_limit] + "\n[Content truncated for AI processing...]"
    
    question_types_str = ", ".join(question_types)
    
    prompt = f"""Based on the following educational content about {topic}, generate EXACTLY {question_count} {difficulty}-level questions.

Content:
{content}

Additional Context: {context}

Generate a mix of question types from: {question_types_str}

Generate questions in this EXACT JSON format (no deviations):
[
    {{
        "q": "Question text here?",
        "answer": "correct answer",
        "keywords": ["alternative1", "alternative2"],
        "feedback": "Educational explanation",
        "type": "short_answer|multiple_choice|true_false",
        "options": ["option1", "option2", "option3", "option4"]
    }}
]

CRITICAL REQUIREMENTS:
- Return EXACTLY {question_count} questions in the array
- Each question must have fields: q, answer, keywords, feedback, type
- For multiple_choice questions: include "options" array with 2-4 choices, answer must match one option exactly
- For true_false questions: answer must be "true" or "false", no options needed
- For short_answer questions: include keywords array for alternative answers, no options needed
- Keep answers concise (under 50 words)
- Keep feedback brief (under 100 words)
- Focus on {difficulty} difficulty level
- Questions should test understanding of {topic}
- Return ONLY the JSON array, no other text
- Do NOT use markdown code blocks
- Ensure valid JSON with proper commas and brackets

Start response with [ and end with ]"""
    
    response = call_ai_api(prompt, max_tokens=5000)
    try:
        # Extract JSON from response (handles markdown code blocks)
        clean_json = extract_json_from_response(response)
        
        # Try to fix incomplete JSON
        fixed_json = fix_incomplete_json(clean_json)
        
        print(f"DEBUG: Original response length: {len(response)}")
        print(f"DEBUG: Extracted JSON length: {len(clean_json)}")
        print(f"DEBUG: Fixed JSON: {fixed_json[:200]}...")
        
        questions_data = json.loads(fixed_json)
        return questions_data
    except json.JSONDecodeError as e:
        # If not valid JSON, return error with more details
        print(f"DEBUG: JSON Parse Error at position {e.pos}: {str(e)}")
        print(f"DEBUG: Problematic JSON section: {response[max(0, e.pos-50):e.pos+50]}")
        return {"error": f"AI returned invalid JSON: {response[:500]}... | Parse error: {str(e)}"}

def grade_answer_with_ai(question, correct_answer, student_answer, confidence_threshold=80):
    """Use AI to grade student answers with semantic understanding"""
    # Check if AI is configured
    if not is_ai_configured():
        return {
            "correct": False,
            "confidence": 0,
            "explanation": "AI grading not available - API key not configured"
        }
    
    prompt = f"""
    Grade this student answer using semantic understanding:
    
    Question: {question}
    Expected Answer: {correct_answer}
    Student Answer: {student_answer}
    
    Evaluate if the student answer is semantically correct, even if worded differently.
    Consider synonyms, alternative phrasings, and equivalent commands/concepts.
    
    Respond in this JSON format:
    {{
        "correct": true/false,
        "confidence": 0-100,
        "explanation": "Brief explanation of your decision"
    }}
    
    Be generous with partial credit for answers that show understanding.
    
    IMPORTANT: Return ONLY the JSON object. Do not use markdown code blocks. Do not include any explanatory text. Just the raw JSON object.
    """
    
    response = call_ai_api(prompt, max_tokens=200)
    
    # Check if response is an error message
    if isinstance(response, str) and ('error' in response.lower() or 'api' in response.lower()):
        return {
            "correct": False,
            "confidence": 0,
            "explanation": f"AI grading error: {response[:100]}"
        }
    
    try:
        # Extract JSON from response (handles markdown code blocks)
        clean_json = extract_json_from_response(response)
        result = json.loads(clean_json)
        # Apply confidence threshold
        if result.get('confidence', 0) < confidence_threshold:
            result['correct'] = False
            result['explanation'] += f" (Below {confidence_threshold}% confidence threshold)"
        return result
    except json.JSONDecodeError as e:
        return {
            "correct": False,
            "confidence": 0,
            "explanation": "Error: Could not parse AI response"
        }

# Helper function for fuzzy answer checking
def check_answer_fuzzy(user_answer, question_data, similarity_threshold=0.8):
    """
    Check user answer against correct answer and alternatives with fuzzy matching.
    Returns tuple: (is_correct, feedback_type, similarity_score)
    """
    question_type = question_data.get('type', 'short_answer')
    user_answer = user_answer.strip()
    correct_answer = question_data.get("answer", "").strip()
    
    # Handle different question types
    if question_type == 'true_false':
        # Normalize true/false answers
        user_normalized = normalize_true_false_answer(user_answer)
        correct_normalized = normalize_true_false_answer(correct_answer)
        
        if user_normalized == correct_normalized:
            return True, "Correct!", 1.0
        else:
            return False, "Incorrect", 0
    
    elif question_type == 'multiple_choice':
        # For multiple choice, check exact match with correct answer or option labels
        user_answer_lower = user_answer.lower().strip()
        correct_answer_lower = correct_answer.lower().strip()
        
        # Check exact match
        if user_answer_lower == correct_answer_lower:
            return True, "Correct choice!", 1.0
        
        # Check if user entered option letter (a, b, c, d)
        options = question_data.get('options', [])
        if len(user_answer) == 1 and user_answer.lower() in 'abcd':
            option_index = ord(user_answer.lower()) - ord('a')
            if 0 <= option_index < len(options):
                if options[option_index].lower().strip() == correct_answer_lower:
                    return True, "Correct choice!", 1.0
        
        # Check if user typed the full option text
        for option in options:
            if option.lower().strip() == user_answer_lower:
                if option.lower().strip() == correct_answer_lower:
                    return True, "Correct choice!", 1.0
        
        return False, "Incorrect choice", 0
    
    else:  # short_answer (default) - use enhanced fuzzy matching
        user_answer = user_answer.lower().strip()
        correct_answer = correct_answer.lower().strip()
        
        # Check exact match with correct answer
        if user_answer == correct_answer:
            return True, "Exact match!", 1.0
        
        # Check exact match with keywords (alternatives)
        keywords = question_data.get("keywords", [])
        if isinstance(keywords, str):
            keywords = [k.strip().lower() for k in keywords.split(',') if k.strip()]
        else:
            keywords = [str(k).strip().lower() for k in keywords]
        
        if user_answer in keywords:
            return True, "Correct alternative!", 1.0
        
        # Enhanced fuzzy matching for short answers
        # 1. Check for partial matches (substring)
        if len(user_answer) > 3:  # Only for longer answers
            if user_answer in correct_answer or correct_answer in user_answer:
                return True, "Partial match!", 0.8
            
            # Check partial matches with keywords
            for keyword in keywords:
                if user_answer in keyword or keyword in user_answer:
                    return True, f"Partial match with '{keyword}'!", 0.8
        
        # 2. Fuzzy matching with correct answer
        correct_similarity = difflib.SequenceMatcher(None, user_answer, correct_answer).ratio()
        if correct_similarity > similarity_threshold:
            return True, "Close enough to correct answer!", correct_similarity
        
        # 3. Fuzzy matching with alternatives
        best_similarity = 0
        best_match = ""
        for keyword in keywords:
            similarity = difflib.SequenceMatcher(None, user_answer, keyword).ratio()
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = keyword
        
        if best_similarity > similarity_threshold:
            return True, f"Close enough to '{best_match}'!", best_similarity
        
        # 4. Word-based fuzzy matching (split into words)
        user_words = set(user_answer.split())
        correct_words = set(correct_answer.split())
        
        if user_words and correct_words:
            word_overlap = len(user_words.intersection(correct_words)) / len(correct_words)
            if word_overlap >= 0.7:  # 70% word overlap
                return True, "Word-based match!", word_overlap
        
        return False, "Incorrect", 0

def normalize_true_false_answer(answer):
    """Normalize true/false answers to handle various inputs"""
    answer_lower = answer.lower().strip()
    
    # True variations
    if answer_lower in ['true', 't', 'yes', 'y', '1', 'correct', 'right']:
        return 'true'
    
    # False variations
    if answer_lower in ['false', 'f', 'no', 'n', '0', 'incorrect', 'wrong']:
        return 'false'
    
    return answer_lower

# Helper function to save leaderboard data
def save_leaderboard(player_name, score, total_time, correct_answers, wrong_answers, game_mode="adventure", level=None):
    print(f"[DEBUG SAVE_LEADERBOARD] Called with: player={player_name}, score={score}, mode={game_mode}, is_student={session.get('is_student')}")
    
    # Save to student leaderboard if it's a logged-in student
    if session.get('is_student') and session.get('student_id'):
        # Get actual student name from students.json instead of relying on player_name
        student_id = session.get('student_id')
        students = load_students()
        student = next((s for s in students if s['id'] == student_id), None)
        
        # Use student's full name if available, otherwise use the passed player_name
        if student and 'full_name' in student:
            player_name = student['full_name']
        elif session.get('student_name'):
            player_name = session.get('student_name')
        
        print(f"Saving student leaderboard data: {player_name}, mode: {game_mode}, level: {level}")
        
        record = {
            "player": player_name,
            "student_id": student_id,
            "score": score,
            "time": round(total_time, 2),
            "correct_answers": correct_answers,
            "wrong_answers": wrong_answers,
            "game_mode": game_mode,
            "date": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # Add level for adventure mode
        if game_mode == "adventure" and level is not None:
            record["level"] = level
        
        try:
            with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
                leaderboard = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            leaderboard = []

        leaderboard.append(record)

        with open(LEADERBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(leaderboard, f, indent=4)
        
        print(f"[DEBUG SAVE_LEADERBOARD] Successfully saved to STUDENT leaderboard: {LEADERBOARD_FILE}")
    
    # Save to guest leaderboard if it's a guest player (not a student)
    else:
        print(f"Saving guest leaderboard data: {player_name}, mode: {game_mode}, level: {level}")
        
        record = {
            "player": player_name,
            "score": score,
            "time": round(total_time, 2),
            "correct_answers": correct_answers,
            "wrong_answers": wrong_answers,
            "game_mode": game_mode,
            "date": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # Add level for adventure mode
        if game_mode == "adventure" and level is not None:
            record["level"] = level
        
        try:
            with open(GUEST_LEADERBOARD_FILE, "r", encoding="utf-8") as f:
                guest_leaderboard = json.load(f)
            print(f"[DEBUG SAVE_LEADERBOARD] Loaded {len(guest_leaderboard)} existing guest records")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"[DEBUG SAVE_LEADERBOARD] Creating new guest leaderboard file: {e}")
            guest_leaderboard = []

        guest_leaderboard.append(record)
        print(f"[DEBUG SAVE_LEADERBOARD] Added record, now have {len(guest_leaderboard)} total records")

        with open(GUEST_LEADERBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(guest_leaderboard, f, indent=4)
        
        print(f"[DEBUG SAVE_LEADERBOARD] Successfully saved to GUEST leaderboard: {GUEST_LEADERBOARD_FILE}")

def load_leaderboard():
    """Load leaderboard data from file"""
    try:
        with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def reset_test_yourself_session():
    """Completely reset Test Yourself mode session data"""
    test_keys = ['test_question_ids', 'test_q_index', 'test_correct', 
                'test_start_time', 'test_time_limit', 'test_user_answers']
    for key in test_keys:
        session.pop(key, None)

def reset_endless_mode_session():
    """Completely reset Endless mode session data"""
    endless_keys = ['endless_score', 'endless_hp', 'endless_streak', 'endless_highest_streak',
                   'endless_total_answered', 'endless_correct', 'endless_wrong', 'endless_start_time',
                   'endless_score_initialized', 'endless_question_start', 'endless_current_question',
                   'endless_feedback_list']
    for key in endless_keys:
        session.pop(key, None)

# Auto-save functionality
def log_analytics_event(event_type, data=None):
    """Log analytics event if analytics are enabled"""
    settings = get_current_game_settings()
    if not settings.get('analytics_enabled', True):
        return
    
    try:
        analytics_data = {
            'timestamp': time.time(),
            'event_type': event_type,
            'player_name': session.get('player_name', 'Anonymous'),
            'session_id': session.get('_id', 'unknown'),
            'data': data or {}
        }
        
        # Append to analytics file
        analytics_file = 'data/analytics.json'
        analytics_log = []
        
        # Load existing analytics
        try:
            with open(analytics_file, 'r', encoding='utf-8') as f:
                analytics_log = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            analytics_log = []
        
        analytics_log.append(analytics_data)
        
        # Keep only last 1000 events to prevent file from growing too large
        if len(analytics_log) > 1000:
            analytics_log = analytics_log[-1000:]
        
        # Save analytics
        os.makedirs('data', exist_ok=True)
        with open(analytics_file, 'w', encoding='utf-8') as f:
            json.dump(analytics_log, f, indent=2)
            
    except Exception as e:
        print(f"Analytics logging failed: {e}")

def log_student_answer(student_id, student_name, question_id, question_text, student_answer, correct_answer, is_correct, game_mode, level=None):
    """Log student answers in real-time and notify teachers via WebSocket"""
    try:
        # Create answer log entry
        answer_log = {
            'timestamp': datetime.now().isoformat(),
            'student_id': student_id,
            'student_name': student_name,
            'question_id': question_id,
            'question_text': question_text[:100] + '...' if len(question_text) > 100 else question_text,
            'student_answer': student_answer,
            'correct_answer': correct_answer,
            'is_correct': is_correct,
            'game_mode': game_mode,
            'level': level
        }
        
        # Load existing answer logs
        try:
            with open('data/student_answers_log.json', 'r') as f:
                answers_data = json.load(f)
        except FileNotFoundError:
            answers_data = []
        
        # Add new answer log
        answers_data.append(answer_log)
        
        # Keep only last 500 answers to prevent file from growing too large
        if len(answers_data) > 500:
            answers_data = answers_data[-500:]
        
        # Save updated answer logs
        os.makedirs('data', exist_ok=True)
        with open('data/student_answers_log.json', 'w') as f:
            json.dump(answers_data, f, indent=2)
        
        # Emit real-time update to teachers (only if socketio is available)
        try:
            socketio.emit('student_answer', answer_log, room='teachers')
        except:
            pass  # SocketIO not available, skip real-time update
        
    except Exception as e:
        print(f"Error logging student answer: {e}")

def auto_save_progress():
    """Auto-save player progress if enabled in settings"""
    settings = get_current_game_settings()
    if not settings.get('auto_save', True):
        return
    
    # Only save if we have meaningful progress to save
    if not session.get('player_name') or not session.get('selected_level'):
        return
    
    try:
        progress_data = {
            'player_name': session.get('player_name', 'Anonymous'),
            'selected_level': session.get('selected_level'),
            'score': session.get('score', 0),
            'player_hp': session.get('player_hp', 100),
            'enemy_hp': session.get('enemy_hp', 50),
            'q_index': session.get('q_index', 0),
            'correct_answers': session.get('correct_answers', 0),
            'wrong_answers': session.get('wrong_answers', 0),
            'highest_unlocked': session.get('highest_unlocked', 1),
            'level_completed': session.get('level_completed', False),
            'character': session.get('character'),
            'timestamp': time.time()
        }
        
        # Save to auto-save file
        auto_save_file = f"data/autosave_{session.get('player_name', 'anonymous')}.json"
        os.makedirs('data', exist_ok=True)
        with open(auto_save_file, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, indent=2)
            
        print(f"Auto-saved progress for {session.get('player_name')}")
        
    except Exception as e:
        print(f"Auto-save failed: {e}")

def load_auto_save_progress(player_name):
    """Load auto-saved progress for a player"""
    try:
        auto_save_file = f"data/autosave_{player_name}.json"
        if not os.path.exists(auto_save_file):
            return None
            
        with open(auto_save_file, 'r', encoding='utf-8') as f:
            progress_data = json.load(f)
            
        # Check if save is recent (within 24 hours)
        if time.time() - progress_data.get('timestamp', 0) > 86400:
            return None
            
        return progress_data
        
    except Exception as e:
        print(f"Failed to load auto-save: {e}")
        return None

# Define the function to get questions for a specific level
def get_questions_for_level(level_number, levels):
    level_info = next((lvl for lvl in levels if lvl["level"] == level_number), None)
    if level_info:
        base_questions = [q for q in questions if q.get("id") in level_info["questions"]]
        
        # Filter by chapter if a specific chapter is selected
        if 'selected_chapter' in session:
            try:
                chapters_data = load_chapters()
                selected_chapter = next((ch for ch in chapters_data.get("chapters", []) 
                                       if ch.get("id") == session['selected_chapter']), None)
                if selected_chapter:
                    chapter_question_ids = set(selected_chapter.get("question_ids", []))
                    base_questions = [q for q in base_questions if q.get("id") in chapter_question_ids]
            except Exception as e:
                print(f"Error filtering by chapter: {e}")
        
        # Apply adaptive difficulty if enabled
        settings = get_current_game_settings()
        if settings.get('adaptive_difficulty', False):
            return apply_adaptive_difficulty(base_questions, level_number)
        return base_questions
    return []

def apply_adaptive_difficulty(base_questions, level_number):
    """Adjust question difficulty based on player performance"""
    # Calculate player's recent accuracy
    correct = session.get('correct_answers', 0)
    wrong = session.get('wrong_answers', 0)
    total = correct + wrong
    
    if total == 0:
        return base_questions  # No performance data yet
    
    accuracy = correct / total
    
    # Adjust difficulty based on performance
    if accuracy > 0.8:  # Player doing very well - increase difficulty
        # Try to get harder questions from higher levels
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                all_levels = json.load(f)
            
            harder_questions = []
            for lvl in all_levels:
                if lvl["level"] > level_number:
                    harder_questions.extend([q for q in questions if q.get("id") in lvl["questions"]])
            
            if harder_questions:
                # Mix 70% base questions with 30% harder questions
                base_count = int(len(base_questions) * 0.7)
                hard_count = min(len(harder_questions), len(base_questions) - base_count)
                return base_questions[:base_count] + random.sample(harder_questions, hard_count)
                
        except Exception:
            pass
            
    elif accuracy < 0.5:  # Player struggling - decrease difficulty
        # Try to get easier questions from lower levels
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                all_levels = json.load(f)
            
            easier_questions = []
            for lvl in all_levels:
                if lvl["level"] < level_number:
                    easier_questions.extend([q for q in questions if q.get("id") in lvl["questions"]])
            
            if easier_questions:
                # Mix 70% base questions with 30% easier questions
                base_count = int(len(base_questions) * 0.7)
                easy_count = min(len(easier_questions), len(base_questions) - base_count)
                return base_questions[:base_count] + random.sample(easier_questions, easy_count)
                
        except Exception:
            pass
    
    return base_questions  # No adjustment needed or possible


# Route for level selection
@app.route('/select_level', methods=['GET', 'POST'])
def select_level():
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except Exception:
        levels = []
    
    # Load chapters for level mode
    chapters_data = load_chapters()
    all_chapters = sorted(chapters_data.get("chapters", []), key=lambda x: x.get("order", 0))
    
    # Group levels by chapters
    chapter_levels = {}
    for chapter in all_chapters:
        chapter_id = chapter.get('id')
        level_range = chapter.get('level_range', [])
        chapter_levels[chapter_id] = {
            'chapter': chapter,
            'levels': [l for l in levels if l.get('level') in level_range]
        }

    # Determine highest unlocked level (simple: highest completed in session, or 1 if none)
    highest_unlocked = session.get('highest_unlocked', 1)
    # Optionally, you could use leaderboard or persistent storage for this

    if request.method == 'POST':
        selected_level = int(request.form.get('level', 1))
        selected_chapter = request.form.get('chapter_id')
        if selected_chapter:
            session['selected_chapter'] = int(selected_chapter)
        session['selected_level'] = selected_level
        # Reset session variables for a new game
        settings = get_current_game_settings()
        session['score'] = 0
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        # Set enemy HP to the current setting
        session['enemy_level'] = selected_level
        session['q_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session["level_start_time"] = time.time()
        session["game_start_time"] = time.time()
        session['level_completed'] = False
        session['enemy_defeated'] = False  # Reset enemy defeated status for new level
        session['current_timer'] = settings.get('question_time_limit', 30)  # Use configurable time limit
        
        # Initialize enemy progression index based on selected level and current progress
        # Only reset to novice if playing level 1 or if no enemy index exists
        try:
            current_enemy_index = session.get('enemy_index')
            
            # If no enemy index exists, start with novice
            # Only reset to novice for level 1 if player hasn't progressed yet
            if current_enemy_index is None or (selected_level == 1 and current_enemy_index == 0):
                novice_idx = 0
                try:
                    with open('data/enemies.json', encoding='utf-8') as ef:
                        enemies_list = json.load(ef)
                    # Look for an enemy by name or level that indicates the novice gnome
                    for i, e in enumerate(enemies_list):
                        name = str(e.get('name', '')).strip().lower()
                        level = e.get('level')
                        if name == 'novice gnome' or level == 1:
                            novice_idx = i
                            break
                except Exception:
                    # If we can't read the file, default to index 0
                    novice_idx = 0
                session['enemy_index'] = int(novice_idx)
            else:
                # Set enemy to match the selected level
                try:
                    with open('data/enemies.json', encoding='utf-8') as ef:
                        enemies_list = json.load(ef)
                    
                    # Set enemy index to match the selected level
                    level_based_index = min(selected_level - 1, len(enemies_list) - 1)
                    session['enemy_index'] = level_based_index
                except Exception:
                    # Keep current enemy index if file can't be read
                    pass
        except Exception:
            session['enemy_index'] = 0
        # Auto-save initial game state
        auto_save_progress()
        
        # Check if student is logged in and has a character
        if session.get('is_student') and 'student_id' in session:
            students = load_students()
            student = next((s for s in students if s['id'] == session['student_id']), None)
            if student and 'selected_character' in student:
                # Student has a character, use it and go to game
                session['character'] = student['selected_character']
                return redirect(url_for('game'))
            else:
                # Student needs to choose a character
                session['redirect_after_character'] = 'game'
                return redirect(url_for('choose_character'))
        else:
            # Non-student user (guest) - always let them choose/change character for each new game
            session['redirect_after_character'] = 'game'
            return redirect(url_for('choose_character'))

    # Pass unlocked info to template (GET request) along with chapters
    return render_template('select_level.html', levels=levels, highest_unlocked=highest_unlocked, 
                         chapters=all_chapters, chapter_levels=chapter_levels)

# Route for the home page
@app.route('/')
def index():
    # Get settings to control feature visibility
    settings = get_current_game_settings()
    # Always redirect to level selection if no level is selected
    # If no selected level, show the merged landing page
    if 'selected_level' not in session:
        return render_template('index.html', settings=settings)
    # If a selected level exists, render the index page
    return render_template('index.html', settings=settings)


# Home route (explicit)
@app.route('/home')
def home():
    return render_template('home.html')


# How-to-play route (ensure it exists)
@app.route('/howto')
def howto():
    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)
    return render_template('howto.html', is_student=is_student)


@app.route('/choose_character', methods=['GET', 'POST'])
def choose_character():
    total_characters = 16
    
    # Check if student is logged in
    if session.get('is_student') and 'student_id' in session:
        student_id = session['student_id']
        students = load_students()
        student = next((s for s in students if s['id'] == student_id), None)
        
        if request.method == 'POST':
            char = request.form.get('character') or request.form.get('character_id')
            try:
                char_id = int(char)
            except (TypeError, ValueError):
                char_id = None
            if char_id and 1 <= char_id <= total_characters:
                # Save character to student profile
                if student:
                    student['selected_character'] = char_id
                    save_students(students)
                session['character'] = char_id
                flash(f'Character {char_id} selected!', 'success')
                # After choosing a character, redirect based on where they came from
                next_url = request.args.get('next') or session.get('redirect_after_character') or 'student_dashboard'
                session.pop('redirect_after_character', None)
                return redirect(url_for(next_url))
        
        # Check if student already has a character
        if student and 'selected_character' in student:
            session['character'] = student['selected_character']
            # Show current character selection
            return render_template('choose_character.html', 
                                 total_characters=total_characters,
                                 current_character=student['selected_character'],
                                 student_name=student['full_name'])
    
    # For non-logged in users or new character selection
    if request.method == 'POST':
        char = request.form.get('character') or request.form.get('character_id')
        try:
            char_id = int(char)
        except (TypeError, ValueError):
            char_id = None
        if char_id and 1 <= char_id <= total_characters:
            session['character'] = char_id
            # After choosing a character, start the game
            return redirect(url_for('game'))
    
    # Check if guest player has a current character to display
    current_char = session.get('character')
    return render_template('choose_character.html', 
                         total_characters=total_characters,
                         current_character=current_char)

# Route for the game page
@app.route('/game', methods=['GET', 'POST'])
def game():
    # Ensure player has selected a level and has basic session data initialized
    if 'selected_level' not in session or 'character' not in session:
        flash('Please select a level first.', 'warning')
        return redirect(url_for('select_level'))
    # Ensure player has a name
    if not session.get('player_name') or not session.get('character'):
        flash('Please set your name first.', 'warning')
        return redirect(url_for('index'))
    
    # Get settings (needed for both initialization and debug checks)
    settings = get_current_game_settings()
    # Initialize HP if not set (for new games)
    if 'player_hp' not in session or 'enemy_hp' not in session:
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        session['score'] = 0
        session['q_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['level_completed'] = False
        session['enemy_defeated'] = False
    
    # Check if the game is over (only after ensuring HP is initialized)
    if session.get('q_index', 0) >= len(questions):
        return redirect(url_for('result'))
    
    # Debug HP check
    current_hp = session.get('player_hp', 100)
    if settings.get('debug_mode', False):
        print(f"DEBUG: HP Check - current_hp: {current_hp}, q_index: {session.get('q_index', 0)}")
    
    if current_hp <= 0:
        if settings.get('debug_mode', False):
            print(f"DEBUG: Redirecting to you_lose due to HP <= 0 (HP: {current_hp})")
        return redirect(url_for('you_lose'))


    # Use selected level from session
    selected_level = session.get('selected_level', 1)
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except FileNotFoundError:
        return "Error: levels.json file not found."
    except json.JSONDecodeError as e:
        return f"Error: Failed to decode levels.json - {e}"

    # Only allow playing the selected level
    current_level = selected_level
    level_questions = get_questions_for_level(current_level, levels)
    print(f"DEBUG: Level {current_level} questions: {len(level_questions)} questions loaded")  # Debugging
    if not level_questions:
        flash(f'No questions available for level {current_level}. Please contact your teacher.', 'error')
        return redirect(url_for('select_level'))

    # Check if we've exceeded max questions
    questions_per_level = settings.get('questions_per_level', 10)
    
    # Mark enemy as defeated if HP <= 0 but continue playing
    if session.get('enemy_hp', BASE_ENEMY_HP) <= 0:
        session['enemy_defeated'] = True
        session['level_completed'] = True
    
    # Only end the level when all questions are answered
    if session['q_index'] >= questions_per_level:
        return redirect(url_for('result'))
    
    # Get question with proper bounds checking
    question_index = session['q_index'] % len(level_questions) if level_questions else 0
    if question_index >= len(level_questions):
        return redirect(url_for('result'))
    question = level_questions[question_index]

    # Always reload enemies from JSON for each game session
    try:
        with open('data/enemies.json', encoding='utf-8') as f:
            enemies = json.load(f)
    except Exception:
        enemies = []
    # Select enemy based on progression index, with fallbacks
    enemy = None
    enemy_index = session.get('enemy_index', 0)
    
    # Try to get enemy by progression index first
    if isinstance(enemies, list) and enemies and 0 <= int(enemy_index) < len(enemies):
        try:
            enemy = enemies[int(enemy_index)]
            print(f"DEBUG: Using enemy from index {enemy_index}: {enemy.get('name', 'Unknown')}")
        except Exception:
            enemy = None

    # Fallback 1: try to find an enemy that matches the current level
    if not enemy:
        enemy = next((e for e in enemies if e.get("level") == current_level), None)
        if enemy:
            print(f"DEBUG: Found level-based enemy for level {current_level}: {enemy.get('name', 'Unknown')}")
    
    # Fallback 2: use enemy based on level progression (level-1 as index)
    if not enemy and isinstance(enemies, list) and enemies:
        level_based_index = min(max(current_level - 1, 0), len(enemies) - 1)
        try:
            enemy = enemies[level_based_index]
            # Update session to match this enemy for consistency
            session['enemy_index'] = level_based_index
            print(f"DEBUG: Using level-based index {level_based_index} for level {current_level}: {enemy.get('name', 'Unknown')}")
        except Exception:
            enemy = None
    
    # Final fallback: default enemy
    if not enemy:
        enemy = {"name": "Unknown Enemy", "avatar": "❓", "taunt": "No enemies found for this level."}
    
    print(f"DEBUG: Final selected enemy for level {current_level}: {enemy.get('name', 'Unknown')} (enemy_index: {session.get('enemy_index')})")

    # Determine enemy image URL to mirror how player avatar images are used
    enemy_image = None
    # If the enemy JSON includes an explicit image field, use it (assume path under static/)
    if isinstance(enemy, dict) and enemy.get('image'):
        enemy_image = url_for('static', filename=enemy.get('image'))
    else:
        # Attempt to find an image in static/enemies matching the enemy name (safe fallback)
        # Create a safe filename from the enemy name (lowercase, replace spaces with underscores)
        if isinstance(enemy, dict) and enemy.get('name'):
            safe_base = ''.join(c if c.isalnum() else '_' for c in enemy.get('name').lower()).strip('_')
            # Check common extensions
            for ext in ['.png', '.jpg', '.jpeg', '.gif']:
                candidate = f'enemies/{safe_base}{ext}'
                candidate_path = os.path.join(os.path.dirname(__file__), 'static', candidate)
                if os.path.exists(candidate_path):
                    enemy_image = url_for('static', filename=candidate)
                    break

            # If not found by safe name, try fuzzy matching against files in static/enemies
            if not enemy_image:
                try:
                    enemy_tokens = [t for t in ''.join(c if c.isalnum() else ' ' for c in enemy.get('name').lower()).split() if t]
                    enemies_dir = os.path.join(os.path.dirname(__file__), 'static', 'enemies')
                    best_match = None
                    best_score = 0
                    if os.path.isdir(enemies_dir):
                        for fn in os.listdir(enemies_dir):
                            fn_lower = fn.lower()
                            name_part = os.path.splitext(fn_lower)[0]
                            file_tokens = [t for t in ''.join(c if c.isalnum() else ' ' for c in name_part).split() if t]
                            # score = number of matching tokens
                            score = sum(1 for t in enemy_tokens if t in file_tokens)
                            # if all tokens match, immediate perfect match
                            if score == len(enemy_tokens) and score > 0:
                                best_match = fn
                                best_score = score
                                break
                            if score > best_score:
                                best_score = score
                                best_match = fn
                        # accept best match if it matches at least one token
                        if best_match and best_score > 0:
                            candidate = f'enemies/{best_match}'
                            enemy_image = url_for('static', filename=candidate)
                except Exception:
                    enemy_image = None

    # Attach enemy_image into the template context

    # Initialize or reset the timer for the current question
    # Always reset timer on GET request
    if request.method == 'GET' or "level_start_time" not in session:
        session["level_start_time"] = time.time()
        # Initialize current_timer if not set
        if "current_timer" not in session:
            session["current_timer"] = settings.get('question_time_limit', 30)

    # Calculate remaining time
    settings = get_current_game_settings()
    elapsed = time.time() - session["level_start_time"]
    time_left = max(0, session.get("current_timer", settings.get('question_time_limit', 30)) - int(elapsed))

    # Debug timer info
    if settings.get('debug_mode', False):
        print(f"DEBUG: Timer - elapsed: {elapsed:.2f}s, time_left: {time_left}s, question_time_limit: {settings.get('question_time_limit', 30)}s, current_timer: {session.get('current_timer', 'not set')}")

    if time_left == 0:
        # Time expired → check timeout behavior setting
        settings = get_current_game_settings()
        session['wrong_answers'] = session.get('wrong_answers', 0) + 1  # Track timeouts as wrong answers
        
        # Check timeout behavior setting
        timeout_behavior = settings.get('timeout_behavior', 'penalty')
        
        if timeout_behavior == 'fail':
            # Immediate fail on timeout
            session['feedback'] = "⏳ Time's up! Timeout results in immediate failure."
            return redirect(url_for('you_lose'))
        else:
            # Apply penalty and continue (default behavior) - use base damage only, not multiplied by level
            timeout_damage = settings['base_damage']
            current_hp = session.get('player_hp', settings['base_player_hp'])
            
            # Ensure HP is properly initialized
            if 'player_hp' not in session:
                session['player_hp'] = settings['base_player_hp']
                current_hp = settings['base_player_hp']
            
            # Apply damage
            session['player_hp'] = current_hp - timeout_damage
            
            # Debug info
            if settings.get('debug_mode', False):
                print(f"DEBUG: Timeout - HP before: {current_hp}, damage: {timeout_damage}, HP after: {session['player_hp']}")
            
            # Check if player has failed due to low HP
            if session['player_hp'] <= 0:
                session['feedback'] = f"⏳ Time's up! You took {timeout_damage} damage and your HP reached 0. Game Over!"
                return redirect(url_for('you_lose'))
            
            session['feedback'] = f"⏳ Time's up! You took {timeout_damage} damage for running out of time. HP: {session['player_hp']}"
            session['q_index'] += 1
            session["level_start_time"] = time.time()  # Reset timer for the next question
            session["current_timer"] = max(10, session.get("current_timer", settings.get('question_time_limit', 30)) - 5)  # Deduct 5s for next question, min 10s
            
            # Auto-save progress after timeout
            auto_save_progress()
            
            return redirect(url_for('feedback'))

    if request.method == 'POST':
        user_answer = request.form.get('answer', '').strip().lower()
        correct_answer = question.get('answer', '').strip().lower()

        # Normalize keywords
        raw_keywords = question.get('keywords', [])
        if isinstance(raw_keywords, str):
            keywords = [k.strip().lower() for k in raw_keywords.split(',') if k.strip()]
        else:
            keywords = [str(k).strip().lower() for k in raw_keywords]

        # Calculate time taken to answer
        time_taken = time.time() - session["level_start_time"]

        # Determine damage and score based on time taken
        if time_taken <= 5:  # First 5 seconds → double damage and double score
            damage = 20
            score = 20
        elif time_taken <= 15:  # Within 15 seconds → regular damage and score
            damage = 10
            score = 10
        else:  # Remaining time → half damage and low score
            damage = 5
            score = 5

        # Only apply this block for the main game mode, not endless or test yourself
        if not session.get('endless_questions') and not session.get('test_questions'):
            # Get current game settings
            settings = get_current_game_settings()
            base_damage = settings['base_damage']
            points_correct = settings['points_correct']
            points_wrong = settings['points_wrong']
            
            # Debug logging if enabled
            if settings.get('debug_mode', False):
                print(f"[DEBUG] Player answer: '{user_answer}' for question: '{question.get('q', 'N/A')[:50]}...'")
                print(f"[DEBUG] Expected answer: '{correct_answer}', Keywords: {keywords}")
                print(f"[DEBUG] Time taken: {time_taken:.2f}s, Damage: {damage}, Score: {score}")
            
            # Get the question feedback
            question_feedback = question.get('feedback', 'No additional information available.')
            
            # Use fuzzy matching for answer checking
            is_correct, feedback_type, similarity_score = check_answer_fuzzy(user_answer, question)
            
            # If student and AI grading is enabled, use AI as fallback for uncertain answers
            if session.get('is_student') and session.get('ai_grading_enabled', False):
                # Use AI grading for short answers with low confidence (< 0.9)
                if question.get('type', 'short_answer') == 'short_answer' and not is_correct and similarity_score < 0.9:
                    try:
                        ai_result = grade_answer_with_ai(
                            question=question.get('q', ''),
                            correct_answer=correct_answer,
                            student_answer=user_answer,
                            confidence_threshold=75
                        )
                        if ai_result.get('correct', False) and ai_result.get('confidence', 0) >= 75:
                            is_correct = True
                            feedback_type = f"AI Grading: {ai_result.get('explanation', 'Accepted')}"
                            similarity_score = ai_result.get('confidence', 0) / 100.0
                    except Exception as e:
                        print(f"AI grading error: {e}")
            
            # Log student answer in real-time
            if 'student_id' in session:
                log_student_answer(
                    student_id=session['student_id'],
                    student_name=session.get('student_name', 'Unknown'),
                    question_id=question.get('id', 'unknown'),
                    question_text=question.get('q', ''),
                    student_answer=user_answer,
                    correct_answer=correct_answer,
                    is_correct=is_correct,
                    game_mode='adventure',
                    level=current_level
                )
            
            # Log answer attempt
            log_analytics_event('answer_submitted', {
                'question_id': question.get('id'),
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'time_taken': time_taken,
                'similarity_score': similarity_score,
                'level': current_level
            })
            
            if is_correct:
                # Calculate speed bonus if enabled
                speed_multiplier = 1.0
                speed_message = ""
                if settings.get('speed_bonus', True) and time_taken <= 10:
                    if time_taken <= 5:
                        speed_multiplier = 2.0
                        speed_message = " (🚀 Lightning bonus: x2 points!)"
                    elif time_taken <= 10:
                        speed_multiplier = 1.5
                        speed_message = " (⚡ Speed bonus: x1.5 points!)"
                
                points_awarded = int(points_correct * speed_multiplier)
                
                # Check if enemy is already defeated for bonus scoring
                enemy_already_defeated = session.get('enemy_defeated', False)
                
                if "Exact answer" in feedback_type or similarity_score == 1.0:
                    # Exact match gets triple damage or bonus points
                    if enemy_already_defeated:
                        # Enemy already defeated - give bonus points instead of damage
                        bonus_points = points_awarded * 2  # Double bonus for exact match after defeat
                        session['score'] += bonus_points
                        session['feedback'] = f"🏆 {feedback_type} Enemy already defeated! Bonus points: {bonus_points} in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                    else:
                        # Normal damage to enemy
                        session['score'] += points_awarded
                        session['enemy_hp'] -= (base_damage * 3)
                        # Ensure enemy HP doesn't go below 0
                        if session['enemy_hp'] <= 0:
                            session['enemy_hp'] = 0
                            session['enemy_defeated'] = True
                        session['feedback'] = f"🔥 {feedback_type} Triple damage: {base_damage*3} and {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                else:
                    # Fuzzy match gets regular damage or normal points
                    similarity_percent = int(similarity_score * 100)
                    if enemy_already_defeated:
                        # Enemy already defeated - give normal points (same as regular scoring)
                        session['score'] += points_awarded
                        session['feedback'] = f"🏆 {feedback_type} ({similarity_percent}% match) Enemy already defeated! Earned {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                    else:
                        # Normal damage to enemy
                        session['score'] += points_awarded
                        session['enemy_hp'] -= base_damage
                        # Ensure enemy HP doesn't go below 0
                        if session['enemy_hp'] <= 0:
                            session['enemy_hp'] = 0
                            session['enemy_defeated'] = True
                        session['feedback'] = f"🎯 {feedback_type} ({similarity_percent}% match) You dealt {base_damage} damage and earned {points_awarded} points in {time_taken:.2f} seconds{speed_message}.<br><br>✅ {question_feedback}"
                
                session["current_timer"] = min(settings.get('question_time_limit', 30) * 2, session.get("current_timer", settings.get('question_time_limit', 30)) + 5)  # Add 5s to timer for next question, max 2x base limit
                session['correct_answers'] = session.get('correct_answers', 0) + 1  # Track correct answers
            else:
                session['player_hp'] -= base_damage  # Use base damage only, not multiplied by level
                session['score'] -= points_wrong  # Deduct points for wrong answer
                session['feedback'] = f"❌ Incorrect! You took {base_damage} damage.<br><br>💡 {question_feedback}"
                session["current_timer"] = max(10, session.get("current_timer", settings.get('question_time_limit', 30)) - 5)  # Deduct 5s from timer for next question, min 10s
                session['wrong_answers'] = session.get('wrong_answers', 0) + 1  # Track wrong answers

        # Auto-save progress after each question
        auto_save_progress()

        # Reset timer for the next question
        session['level_start_time'] = time.time()

        # Move to the next question
        session['q_index'] += 1
        return redirect(url_for('feedback'))

    # Get current settings for display options
    settings = get_current_game_settings()
    
    # Generate dynamic taunt based on the current question
    enemy_taunt = generate_enemy_taunt(question, enemy.get('name', 'Unknown Enemy'))
    
    return render_template('game.html',
                           question=question,
                           score=session['score'],
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           enemy_defeated=session.get('enemy_defeated', False),
                           q_number=question_index + 1,
                           total=questions_per_level,  # Use configurable questions per level
                           level=current_level,
                           enemy=enemy,
                           enemy_image=enemy_image,
                           enemy_taunt=enemy_taunt,
                           time_left=time_left,
                           settings=settings)

# Route for the feedback page
@app.route('/feedback')
def feedback():
    # If player's HP reached 0, go straight to defeat
    if session.get('player_hp', 0) <= 0:
        return redirect(url_for('you_lose'))
    if not session.get('feedback'):
        return redirect(url_for('game'))  # Redirect to the game page if no feedback
    feedback = session['feedback']
    session['feedback'] = None  # Clear feedback after rendering
    return render_template('feedback.html',
                           feedback=feedback,
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           score=session['score'],
                           level=session['enemy_level'])

# Route for the result page
@app.route('/result')
def result():
    # Calculate the final score
    settings = get_current_game_settings()
    bonus = settings['level_bonus'] if session.get('level_completed', False) and session['player_hp'] > 0 else 0
    final_score = session['score'] + bonus

    # Save the player's performance to the leaderboard
    total_time = time.time() - session.get("game_start_time", time.time())
    save_leaderboard(
        player_name=session.get("player_name", "Anonymous"),
        score=final_score,
        total_time=total_time,
        correct_answers=session.get('correct_answers', 0),
        wrong_answers=session.get('wrong_answers', 0),
        game_mode="adventure",
        level=session.get('selected_level', 1)
    )
    
    # Save student progress if logged in
    if session.get('is_student') and session.get('student_id'):
        student_id = session.get('student_id')
        level = session.get('selected_level')
        correct_answers = session.get('correct_answers', 0)
        wrong_answers = session.get('wrong_answers', 0)
        total_questions = correct_answers + wrong_answers
        
        update_student_progress(
            student_id=student_id,
            game_type='adventure_mode',
            level=level,
            score=final_score,
            correct_answers=correct_answers,
            total_questions=total_questions,
            time_taken=total_time
        )

    # If the player died, show lose page
    if session.get('player_hp', 0) <= 0:
        return redirect(url_for('you_lose'))

    # Log level completion
    log_analytics_event('level_completed', {
        'level': session.get('selected_level', 1),
        'final_score': final_score,
        'player_hp': session.get('player_hp', 0),
        'enemy_hp': session.get('enemy_hp', 0),
        'total_time': total_time,
        'correct_answers': session.get('correct_answers', 0),
        'wrong_answers': session.get('wrong_answers', 0),
        'level_completed': session.get('level_completed', False)
    })
    
    # Check if player completed level 10 with 7+ correct answers - show credits
    current_level = session.get('selected_level', 1)
    correct_answers = session.get('correct_answers', 0)
    if current_level == 10 and correct_answers >= 7 and session.get('level_completed', False):
        return redirect(url_for('credits'))

    # If the level is completed, allow to select next level
    next_level = session.get('selected_level', 1) + 1
    try:
        with open("data/levels.json", "r", encoding="utf-8") as f:
            levels = json.load(f)
    except Exception:
        levels = []
    max_level = max([lvl['level'] for lvl in levels], default=1)
    
    # Calculate required accuracy based on settings
    settings = get_current_game_settings()
    total_questions = session.get('correct_answers', 0) + session.get('wrong_answers', 0)
    current_accuracy = (session.get('correct_answers', 0) / total_questions * 100) if total_questions > 0 else 0
    required_accuracy = settings.get('min_accuracy', 70)
    
    can_advance = (
        session.get('level_completed', False)
        and current_accuracy >= required_accuracy
        and next_level <= max_level
    )

    # Unlock the next level if completed AND meets accuracy requirement
    if session.get('level_completed', False):
        prev_highest = session.get('highest_unlocked', 1)
        
        # Only unlock next level if accuracy requirement is met
        if current_accuracy >= required_accuracy and next_level > prev_highest:
            session['highest_unlocked'] = next_level
        elif current_accuracy < required_accuracy:
            # Player completed level but didn't meet accuracy requirement
            # They can replay the same level but can't advance
            pass

        # Auto-save progress after level completion
        auto_save_progress()

        # Advance the enemy_index so a new enemy appears after each completed level
        try:
            with open('data/enemies.json', encoding='utf-8') as f:
                enemies_list = json.load(f)
        except Exception:
            enemies_list = []
        
        if isinstance(enemies_list, list) and enemies_list:
            current_idx = session.get('enemy_index', 0) or 0
            try:
                current_idx = int(current_idx)
            except Exception:
                current_idx = 0
            
            # Advance to next enemy, but don't exceed the list bounds
            next_idx = min(current_idx + 1, len(enemies_list) - 1)
            
            # Also ensure the enemy progression matches or exceeds the level progression
            # Each completed level should advance at least to that level's enemy
            level_based_idx = min(next_level - 1, len(enemies_list) - 1) if next_level <= max_level else len(enemies_list) - 1
            
            # Use the higher of the two indices to ensure proper progression
            final_idx = max(next_idx, level_based_idx)
            session['enemy_index'] = final_idx
            
            print(f"DEBUG: Enemy progression - Level: {next_level}, Current idx: {current_idx}, Next idx: {next_idx}, Level-based idx: {level_based_idx}, Final idx: {final_idx}")
        # If we've unlocked past the maximum level, the player beat the game
        try:
            with open("data/levels.json", "r", encoding="utf-8") as f:
                levels = json.load(f)
        except Exception:
            levels = []
        max_level = max([lvl['level'] for lvl in levels], default=1)
        if session.get('highest_unlocked', 1) > max_level:
            # Player has unlocked beyond max level -> they completed all levels
            return redirect(url_for('you_win'))

    return render_template('result.html', score=final_score,
                           player_hp=session['player_hp'],
                           enemy_hp=session['enemy_hp'],
                           outcome="Game Over",
                           level=session['enemy_level'],
                           enemy={"name": "Unknown Enemy", "avatar": "❓"},
                           can_advance=can_advance,
                           next_level=next_level)

# Route for the search functionality
from whoosh.qparser import QueryParser

@app.route('/search', methods=['GET'])
def search():
    query_text = request.args.get('q', '').strip()
    results = []
    
    if query_text:
        query_lower = query_text.lower()
        # Search through questions using multiple criteria
        for question in questions:
            match_found = False
            
            # Search in question text
            if query_lower in question.get('q', '').lower():
                match_found = True
            
            # Search in answer
            elif query_lower in question.get('answer', '').lower():
                match_found = True
            
            # Search in keywords
            elif question.get('keywords'):
                keywords = question['keywords']
                if isinstance(keywords, str):
                    if query_lower in keywords.lower():
                        match_found = True
                elif isinstance(keywords, list):
                    if any(query_lower in keyword.lower() for keyword in keywords):
                        match_found = True
            # Search in feedback/explanation
            elif question.get('feedback') and query_lower in question.get('feedback', '').lower():
                match_found = True
            
            if match_found:
                results.append(question)
    
    return render_template('search.html', results=results)

# Route for the leaderboard page
@app.route('/leaderboard')
def leaderboard():
    try:
        with open(LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            all_data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        all_data = []

    # Separate leaderboards by game mode
    adventure_data = [entry for entry in all_data if entry.get("game_mode", "adventure") == "adventure"]
    test_yourself_data = [entry for entry in all_data if entry.get("game_mode") == "test_yourself"]
    endless_data = [entry for entry in all_data if entry.get("game_mode") == "endless"]
    
    # Get best score per player for each mode
    def get_best_scores(data, limit=50):
        player_best = {}
        for entry in data:
            player_name = entry.get("player", "Anonymous")
            if player_name not in player_best:
                player_best[player_name] = entry
            else:
                # Keep entry with higher score, or if same score, lower time
                current_best = player_best[player_name]
                if (entry["score"] > current_best["score"] or 
                    (entry["score"] == current_best["score"] and entry["time"] < current_best["time"])):
                    player_best[player_name] = entry
        # Sort by score (highest first), then by time (lowest first)
        sorted_players = sorted(player_best.values(), key=lambda x: (-x["score"], x["time"]))
        return sorted_players[:limit]
    
    # For adventure mode, create per-level leaderboards
    adventure_levels = {}
    for level_num in range(1, 11):  # Levels 1-10
        level_data = [entry for entry in adventure_data if entry.get("level") == level_num]
        adventure_levels[level_num] = get_best_scores(level_data, 10)  # Top 10 per level
    
    # Overall adventure leaderboard (all levels combined)
    adventure_leaderboard = get_best_scores(adventure_data, 50)
    
    test_yourself_leaderboard = get_best_scores(test_yourself_data, 50)
    endless_leaderboard = get_best_scores(endless_data, 50)

    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)

    return render_template("leaderboard.html", 
                         adventure_leaderboard=adventure_leaderboard,
                         adventure_levels=adventure_levels,
                         test_yourself_leaderboard=test_yourself_leaderboard,
                         endless_leaderboard=endless_leaderboard,
                         is_student=is_student)


@app.route('/guest_leaderboard')
def guest_leaderboard():
    try:
        with open(GUEST_LEADERBOARD_FILE, "r", encoding="utf-8") as f:
            all_data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        all_data = []

    # Separate guest leaderboards by game mode
    adventure_data = [entry for entry in all_data if entry.get("game_mode", "adventure") == "adventure"]
    test_yourself_data = [entry for entry in all_data if entry.get("game_mode") == "test_yourself"]
    endless_data = [entry for entry in all_data if entry.get("game_mode") == "endless"]
    
    # Get best score per player for each mode
    def get_best_scores(data, limit=50):
        player_best = {}
        for entry in data:
            player_name = entry.get("player", "Anonymous")
            if player_name not in player_best:
                player_best[player_name] = entry
            else:
                # Keep entry with higher score, or if same score, lower time
                current_best = player_best[player_name]
                if (entry["score"] > current_best["score"] or 
                    (entry["score"] == current_best["score"] and entry["time"] < current_best["time"])):
                    player_best[player_name] = entry
        # Sort by score (highest first), then by time (lowest first)
        sorted_players = sorted(player_best.values(), key=lambda x: (-x["score"], x["time"]))
        return sorted_players[:limit]
    
    # For adventure mode, create per-level leaderboards
    adventure_levels = {}
    for level_num in range(1, 11):  # Levels 1-10
        level_data = [entry for entry in adventure_data if entry.get("level") == level_num]
        adventure_levels[level_num] = get_best_scores(level_data, 10)  # Top 10 per level
    
    # Overall adventure leaderboard (all levels combined)
    adventure_leaderboard = get_best_scores(adventure_data, 50)
    
    test_yourself_leaderboard = get_best_scores(test_yourself_data, 50)
    endless_leaderboard = get_best_scores(endless_data, 50)

    return render_template("guest_leaderboard.html", 
                         adventure_leaderboard=adventure_leaderboard,
                         adventure_levels=adventure_levels,
                         test_yourself_leaderboard=test_yourself_leaderboard,
                         endless_leaderboard=endless_leaderboard)


@app.route('/you_win')
def you_win():
    # Simple win page when all levels are completed
    # Reset enemy state to original when the player finishes the game
    try:
        settings = get_current_game_settings()
        session['enemy_index'] = 0
        session['enemy_hp'] = settings['base_enemy_hp']
        session.pop('enemy_level', None)
    except Exception:
        pass
    
    # Check if student is logged in to provide proper navigation context
    is_student = session.get('is_student', False)
    return render_template('you_win.html', is_student=is_student)


@app.route('/credits')
def credits():
    """Display credits page after completing level 10 with 7+ correct answers"""
    # Get final score and stats from session
    settings = get_current_game_settings()
    bonus = settings['level_bonus'] if session.get('level_completed', False) and session['player_hp'] > 0 else 0
    final_score = session.get('score', 0) + bonus
    correct_answers = session.get('correct_answers', 0)
    wrong_answers = session.get('wrong_answers', 0)
    total_time = time.time() - session.get("game_start_time", time.time())
    
    # Check if student is logged in
    is_student = session.get('is_student', False)
    
    return render_template('credits.html', 
                         final_score=final_score,
                         correct_answers=correct_answers,
                         wrong_answers=wrong_answers,
                         total_time=round(total_time, 2),
                         is_student=is_student)


@app.route('/you_lose')
def you_lose():
    # Reset only the current level progress, not the entire game
    try:
        settings = get_current_game_settings()
        current_level = session.get('selected_level', 1)
        
        # Reset level-specific progress but keep level selection
        session['q_index'] = 0
        session['player_hp'] = settings['base_player_hp']
        session['enemy_hp'] = settings['base_enemy_hp']
        session['enemy_index'] = 0
        session['correct_answers'] = 0
        session['wrong_answers'] = 0
        session['score'] = 0
        session['level_completed'] = False
        session['enemy_defeated'] = False
        session.pop('feedback', None)
        
        # Keep the selected level so they can retry the same level
        session['selected_level'] = current_level
        session['enemy_level'] = current_level
        
        # Log the level failure for analytics
        log_analytics_event('level_failed', {
            'level': current_level,
            'reason': 'hp_zero'
        })
        
    except Exception as e:
        print(f"Error in you_lose route: {e}")
    
    return render_template('you_lose.html')

@app.route('/quit_game')
def quit_game():
    """Handle player quitting mid-game"""
    try:
        current_level = session.get('selected_level', 1)
        current_score = session.get('score', 0)
        questions_answered = session.get('q_index', 0)
        
        # Log the quit action for analytics
        log_analytics_event('game_quit', {
            'level': current_level,
            'score': current_score,
            'questions_answered': questions_answered,
            'reason': 'player_quit'
        })
        
        # Clear current game progress but keep player info
        game_progress_keys = ['q_index', 'player_hp', 'enemy_hp', 'score', 
                             'correct_answers', 'wrong_answers', 'level_completed', 
                             'enemy_defeated', 'feedback', 'current_timer']
        
        for key in game_progress_keys:
            session.pop(key, None)
        
        flash(f'Game quit. You scored {current_score} points on Level {current_level}.', 'info')
        
    except Exception as e:
        print(f"Error in quit_game route: {e}")
        flash('Game ended.', 'info')
    
    return redirect(url_for('select_level'))

@app.route('/quit_test_yourself')
def quit_test_yourself():
    """Handle player quitting Test Yourself mode - auto-save to leaderboard"""
    try:
        questions_answered = session.get('test_q_index', 0)
        correct_answers = session.get('test_correct', 0)
        total_questions = len(session.get('test_question_ids', []))
        test_start_time = session.get('test_start_time', time.time())
        time_taken = time.time() - test_start_time
        score = correct_answers * 10  # Same scoring as normal completion
        player_name = session.get('player_name', 'Anonymous')
        
        # Auto-save to leaderboard when quitting
        save_leaderboard(
            player_name=player_name,
            score=score,
            total_time=time_taken,
            correct_answers=correct_answers,
            wrong_answers=questions_answered - correct_answers,
            game_mode="test_yourself"
        )
        
        # Save student progress if logged in
        if session.get('is_student') and session.get('student_id'):
            student_id = session.get('student_id')
            update_student_progress(
                student_id=student_id,
                game_type='test_yourself',
                level=None,
                score=score,
                correct_answers=correct_answers,
                total_questions=total,
                time_taken=time_taken
            )
        
        # Log the quit action for analytics
        log_analytics_event('test_yourself_quit', {
            'questions_answered': questions_answered,
            'correct_answers': correct_answers,
            'score': score,
            'reason': 'player_quit'
        })
        
        # Clear all test yourself session data completely
        reset_test_yourself_session()
        
        accuracy = (correct_answers / questions_answered * 100) if questions_answered > 0 else 0
        flash(f'Test Yourself completed (quit). Score: {score}, Accuracy: {accuracy:.1f}% - Saved to leaderboard!', 'success')
        
    except Exception as e:
        print(f"Error in quit_test_yourself route: {e}")
        flash('Test session ended.', 'info')
    
    return redirect(url_for('leaderboard'))

@app.route('/quit_endless')
def quit_endless():
    """Handle player quitting Endless mode - auto-save to leaderboard"""
    try:
        total_answered = session.get('endless_total_answered', 0)
        correct_answers = session.get('endless_correct', 0)
        wrong_answers = session.get('endless_wrong', 0)
        final_score = session.get('endless_score', 0)
        highest_streak = session.get('endless_highest_streak', 0)
        endless_start_time = session.get('endless_start_time', time.time())
        total_time = time.time() - endless_start_time
        player_name = session.get('player_name', 'Anonymous')
        
        # Auto-save to leaderboard when quitting
        save_leaderboard(
            player_name=player_name,
            score=final_score,
            total_time=total_time,
            correct_answers=correct_answers,
            wrong_answers=wrong_answers,
            game_mode="endless"
        )
        
        # Save student progress if logged in
        if session.get('is_student') and session.get('student_id'):
            student_id = session.get('student_id')
            update_student_progress(
                student_id=student_id,
                game_type='endless_mode',
                level=None,
                score=final_score,
                correct_answers=correct_answers,
                total_questions=total_answered,
                time_taken=total_time
            )
        
        # Log the quit action for analytics
        log_analytics_event('endless_mode_quit', {
            'questions_answered': total_answered,
            'final_score': final_score,
            'max_streak': highest_streak,
            'reason': 'player_quit'
        })
        
        # Clear all endless mode session data completely
        reset_endless_mode_session()
        
        flash(f'Endless Mode completed (quit). Score: {final_score}, {total_answered} questions - Saved to leaderboard!', 'success')
        
    except Exception as e:
        print(f"Error in quit_endless route: {e}")
        flash('Endless session ended.', 'info')
    
    return redirect(url_for('leaderboard'))

# ------------------- TEST YOURSELF MODE -------------------
import random

@app.route('/test_yourself', methods=['GET', 'POST'])
def test_yourself():
    # Check if Test Yourself mode is enabled
    settings = get_current_game_settings()
    if not settings.get('test_yourself_enabled', True):
        flash('Test Yourself mode is currently disabled.', 'error')
        return redirect(url_for('index'))
    
    # Reset test state for a true new start (GET with ?new=1) or if no session data exists
    if (request.method == 'GET' and request.args.get('new') == '1') or not session.get('test_question_ids'):
        # Completely reset session to ensure clean start
        reset_test_yourself_session()
        session['test_user_answers'] = []
        # ...existing code...
        
        # Use questions from test_yourself pool
        test_pool_questions = get_questions_for_pool('test_yourself')
        if not test_pool_questions:
            test_pool_questions = questions  # Fallback to all questions
        
        valid_questions = [q for q in test_pool_questions if q.get('q') and str(q.get('q')).strip()]
        if not valid_questions:
            session['test_question_ids'] = []
        elif len(valid_questions) >= 40:
            # Use random.sample to guarantee no duplicates (returns unique selection)
            selected = random.sample(valid_questions, 40)
            # Extract IDs and convert to set for guaranteed uniqueness
            question_ids_set = set(q['id'] for q in selected)
            # Convert back to list and shuffle
            question_ids_list = list(question_ids_set)
            random.shuffle(question_ids_list)
            session['test_question_ids'] = question_ids_list
        else:
            # If fewer than 40 questions, use all available without repeats
            # Use set to ensure absolute uniqueness even with small pools
            unique_questions = list({q['id']: q for q in valid_questions}.values())
            random.shuffle(unique_questions)
            session['test_question_ids'] = [q['id'] for q in unique_questions]
        
        # Debug: Verify uniqueness using set comparison
        question_ids = session['test_question_ids']
        unique_ids = set(question_ids)
        print(f"[DEBUG TEST INIT] Selected {len(question_ids)} questions, {len(unique_ids)} unique IDs (set-verified)")
        if len(question_ids) != len(unique_ids):
            duplicate_ids = [id for id in unique_ids if question_ids.count(id) > 1]
            print(f"[CRITICAL TEST INIT] Duplicate question IDs found despite set conversion: {duplicate_ids}")
            # Force fix by using only unique IDs
            session['test_question_ids'] = list(unique_ids)
            random.shuffle(session['test_question_ids'])
        
        session['test_q_index'] = 0
        session['test_correct'] = 0
        session['test_start_time'] = time.time()
        session['test_time_limit'] = 60 * 60  # 1 hour in seconds

    # Calculate timer
    total_seconds_left = max(0, int(session.get('test_time_limit', 3600) - (time.time() - session.get('test_start_time', time.time()))))
    time_left_min = total_seconds_left // 60
    time_left_sec = total_seconds_left % 60
    q_index = session.get('test_q_index', 0)
    test_question_ids = session.get('test_question_ids', [])
    
    # Debug: Print current state
    print(f"[DEBUG TEST] GET request - q_index={q_index}, total_test_questions={len(test_question_ids)}")
    if q_index < len(test_question_ids):
        print(f"[DEBUG TEST] Current question ID: {test_question_ids[q_index]}")
    
    # Rebuild the test_questions list from global questions using IDs
    if not questions:
        flash('No questions available. Please contact your teacher.', 'error')
        return redirect(url_for('index'))
    
    try:
        id_to_question = {q['id']: q for q in questions}
        test_questions = [id_to_question[qid] for qid in test_question_ids if qid in id_to_question]
        
        # Store only question count, not full IDs list to reduce session size
        if 'test_total_questions' not in session:
            session['test_total_questions'] = len(test_question_ids)
    except Exception as e:
        print(f"[ERROR] Failed to rebuild test_questions: {e}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))
    
    # Check if we've reached the end or time is up (40 questions = indices 0-39)
    # Using > 39 instead of >= 40 to ensure question 40 (index 39) is displayed
    if not test_questions or q_index > 39 or total_seconds_left <= 0:
        print(f"[DEBUG] REDIRECT TO RESULT: test_q_index={q_index}, test_questions={len(test_questions)}, total_seconds_left={total_seconds_left}, test_user_answers={len(session.get('test_user_answers', []))}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))

    # Skip invalid questions
    while q_index < len(test_questions) and q_index < 40:
        try:
            if test_questions[q_index].get('q') and str(test_questions[q_index].get('q')).strip():
                break
            q_index += 1
            session['test_q_index'] = q_index
        except (IndexError, KeyError):
            q_index += 1
            session['test_q_index'] = q_index
    
    # Final check after skipping invalid questions
    if q_index > 39:
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))
    
    # Safety check for question existence
    try:
        question = test_questions[q_index]
        if not question or not question.get('q') or not str(question.get('q')).strip():
            raise IndexError("Invalid question")
    except (IndexError, KeyError) as e:
        print(f"[ERROR] Question access error at index {q_index}: {e}")
        session['test_q_index'] = 40
        return redirect(url_for('test_yourself_result'))

    correct_count = session.get('test_correct', 0)
    if request.method == 'POST':
        try:
            user_answer = request.form.get('answer', '').strip().lower()
            correct_answer = question.get('answer', '').strip().lower()
            # Normalize keywords
            raw_keywords = question.get('keywords', [])
            if isinstance(raw_keywords, str):
                keywords = [k.strip().lower() for k in raw_keywords.split(',') if k.strip()]
            else:
                keywords = [str(k).strip().lower() for k in raw_keywords]
            # Use fuzzy matching for test mode
            is_correct, feedback_type, similarity_score = check_answer_fuzzy(user_answer, question)
            
            # If student and AI grading is enabled, use AI as fallback for uncertain answers
            if session.get('is_student') and session.get('ai_grading_enabled', False):
                # Use AI grading for short answers with low confidence (< 0.9)
                if question.get('type', 'short_answer') == 'short_answer' and not is_correct and similarity_score < 0.9:
                    try:
                        ai_result = grade_answer_with_ai(
                            question=question.get('q', ''),
                            correct_answer=correct_answer,
                            student_answer=user_answer,
                            confidence_threshold=75
                        )
                        if ai_result.get('correct', False) and ai_result.get('confidence', 0) >= 75:
                            is_correct = True
                            feedback_type = f"AI Grading: {ai_result.get('explanation', 'Accepted')}"
                            similarity_score = ai_result.get('confidence', 0) / 100.0
                    except Exception as e:
                        print(f"AI grading error: {e}")
            
            # Log student answer in real-time
            if 'student_id' in session:
                log_student_answer(
                    student_id=session['student_id'],
                    student_name=session.get('student_name', 'Unknown'),
                    question_id=question.get('id', 'unknown'),
                    question_text=question.get('q', ''),
                    student_answer=user_answer,
                    correct_answer=correct_answer,
                    is_correct=is_correct,
                    game_mode='test_yourself'
                )
            
            session['test_user_answers'].append({
                'question': question.get('q', '')[:150],  # Further truncate questions
                'user_answer': user_answer[:100],  # Truncate user answers
                'correct_answer': correct_answer[:100],  # Truncate correct answers
                'correct': is_correct,
                'feedback': question.get('feedback', '')[:150],  # Reduce feedback size
                'match_type': feedback_type[:50] if isinstance(feedback_type, str) else str(feedback_type)[:50],
                'similarity': round(similarity_score, 2) if similarity_score else 0
            })
            # Keep only essential answers, limit to 40
            if len(session['test_user_answers']) > 40:
                session['test_user_answers'] = session['test_user_answers'][-40:]
            if is_correct:
                session['test_correct'] = correct_count + 1
            
            # Increment question index
            new_q_index = q_index + 1
            session['test_q_index'] = new_q_index
            
            # Debug logging
            print(f"[DEBUG TEST POST] Answered Q{q_index + 1} (ID={question.get('id')}), correct={is_correct}")
            print(f"[DEBUG TEST POST] Moving to next: new_q_index={new_q_index}, total_questions={len(test_question_ids)}")
            if new_q_index < len(test_question_ids):
                print(f"[DEBUG TEST POST] Next question ID will be: {test_question_ids[new_q_index]}")
            
            return redirect(url_for('test_yourself'))
        except Exception as e:
            print(f"[ERROR] Test yourself mode POST error: {str(e)}")
            # Force game over on error to prevent crash
            session['test_q_index'] = 40
            return redirect(url_for('test_yourself_result'))

    return render_template('test_yourself.html',
                          question=question,
                          q_number=q_index + 1,
                          q_index=q_index,
                          test_questions=test_questions,
                          correct_count=correct_count,
                          time_left_min=time_left_min,
                          time_left_sec=time_left_sec,
                          total_seconds_left=total_seconds_left)
